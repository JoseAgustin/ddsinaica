#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
informe_dicotomico.py
=====================
Calcula estadísticos dicotómicos mensuales (POD, FAR, CSI, TSS, PC, BIAS)
a partir de los CSV diarios de `combinado/ajustados/` y genera un documento
Word (.docx) con los resultados por contaminante, ciudad y horizonte.

Contaminantes y umbrales — NOM-172-SEMARNAT-2023 (DOF 25/01/2024):
    Categoría "Mala"     (riesgo Alto / naranja):
        O3  = 135 ppb  | PM10 = 132 µg/m³ | PM2.5 = 79 µg/m³  | SO2 = 185 ppb
    Categoría "Muy Mala" (riesgo Muy Alto / rojo):
        O3  = 175 ppb  | PM10 = 213 µg/m³ | PM2.5 = 130 µg/m³ | SO2 = 304 ppb

    La categoría de evaluación se selecciona con --categoria (default: mala).

Tabla de contingencia 2×2 (evento = valor ≥ umbral normativo):
    H  — acierto          (obs EVENTO,    mod EVENTO)
    M  — fallo            (obs EVENTO,    mod NO-evento)
    F  — falsa alarma     (obs NO-evento, mod EVENTO)
    C  — rechazo correcto (obs NO-evento, mod NO-evento)

Métricas:
    POD  = H / (H + M)        Prob. de detección             ideal → 1
    FAR  = F / (H + F)        Tasa de falsas alarmas         ideal → 0
    CSI  = H / (H + M + F)    Índice de éxito crítico        ideal → 1
    TSS  = H/(H+M) - F/(F+C)  Pierce Skill Score             ideal → 1
    PC   = (H + C) / N        Porcentaje correcto            ideal → 1
    BIAS = (H + F) / (H + M)  Sesgo de frecuencia            ideal = 1

Control de calidad de datos observados (datos sin curar de SINAICA):
    Los datos de SINAICA se usan en modo crudo. Antes de calcular la tabla
    de contingencia se aplican dos filtros en cascada:

    1. Límites físicos por contaminante (LIMITES_VALIDOS):
       elimina pares donde obs < min_obs, obs > max_obs,
       mod < min_mod o mod > max_mod.
       Captura valores negativos e implausiblemente altos en ambas series.

    2. Filtro IQR sobre observaciones de PM2.5 (CONTAMINANTES_FILTRO_IQR):
       elimina pares donde obs < Q1 - k·IQR  o  obs > Q3 + k·IQR,
       con k = IQR_FACTOR (default 3.0, configurable con --iqr-factor).

Uso:
    python3 informe_dicotomico.py
    python3 informe_dicotomico.py --entrada combinado/ajustados --salida informe.docx
    python3 informe_dicotomico.py --mes 2026-06
    python3 informe_dicotomico.py --ciudades Pachuca Tula CDMX
    python3 informe_dicotomico.py --umbral-pm25 120 --iqr-factor 3.0
    python3 informe_dicotomico.py --categoria mala
    python3 informe_dicotomico.py --categoria muy_mala
    python3 informe_dicotomico.py --help

Dependencias:
    pip install pandas numpy python-docx

Autor  : Pipeline ddsinaica / WRF-Chem — ICAyCC, UNAM
Versión: 1.3.0 (2026-07) — umbrales NOM-172-SEMARNAT-2023; argumento --categoria
"""

import argparse
import glob
import os
import re
import sys
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips

# ──────────────────────────────────────────────────────────────────────────────
# CONFIGURACIÓN GLOBAL
# ──────────────────────────────────────────────────────────────────────────────

# ── Categorías de calidad del aire — NOM-172-SEMARNAT-2023 ───────────────────
#
# Tabla de umbrales por categoría (límite superior del intervalo, inclusive).
# El cálculo dicotómico compara el máximo diario contra el umbral seleccionado.
#
# Categoría "mala"     → riesgo ALTO    (naranja) — se espera que el modelo detecte
# Categoría "muy_mala" → riesgo MUY ALTO (rojo)   — episodios más graves
#
# Fuente: NOM-172-SEMARNAT-2023, Tablas 4–8 (DOF 25/01/2024)
#   O3   : promedio de 1 hora (ppm → ppb × 1000)
#   PM10 : promedio móvil ponderado de 12 h (µg/m³)
#   PM2.5: promedio móvil ponderado de 12 h (µg/m³)
#   SO2  : promedio de 1 hora (ppm → ppb × 1000)
#
CATEGORIAS_NOM172 = {
    "mala": {
        #  límite SUPERIOR de la banda "Mala" (riesgo Alto / naranja)
        "O3":   135.0,   # ppb   (0.135 ppm × 1000)
        "PM10": 132.0,   # µg/m³
        "PM25":  79.0,   # µg/m³
        "SO2":  185.0,   # ppb   (0.185 ppm × 1000)
    },
    "muy_mala": {
        #  límite SUPERIOR de la banda "Muy Mala" (riesgo Muy Alto / rojo)
        "O3":   175.0,   # ppb   (0.175 ppm × 1000)
        "PM10": 213.0,   # µg/m³
        "PM25": 130.0,   # µg/m³
        "SO2":  304.0,   # ppb   (0.304 ppm × 1000)
    },
}

# Categoría activa por defecto — puede cambiarse con --categoria
CATEGORIA_DEFAULT = "mala"

# Alias amigables para argparse
ALIAS_CATEGORIA = {
    "mala":     "mala",
    "malo":     "mala",
    "alto":     "mala",
    "naranja":  "mala",
    "muy_mala": "muy_mala",
    "muymala":  "muy_mala",
    "muy mala": "muy_mala",
    "muy_alto": "muy_mala",
    "muyalto":  "muy_mala",
    "rojo":     "muy_mala",
}

# Variable global de trabajo; se sobreescribe en main() según --categoria
UMBRALES: Dict[str, float] = dict(CATEGORIAS_NOM172[CATEGORIA_DEFAULT])

META_CONT = {
    "O3": {
        "nombre": "Ozono (O\u2083)",
        "unidad": "ppbv",
        "norma":  "NOM-172-SEMARNAT-2023 / NOM-020-SSA1-2021",
        "desc": (
            "Gas oxidante de la troposfera, formado fotoquímicamente a partir de "
            "NO\u2093 y compuestos orgánicos volátiles (COV). Causa irritación "
            "respiratoria y daño a cultivos. "
            "Umbrales NOM-172-SEMARNAT-2023 (promedio 1 h): "
            "Mala = 135 ppb | Muy Mala = 175 ppb."
        ),
    },
    "PM10": {
        "nombre": "Partículas suspendidas gruesas (PM10)",
        "unidad": "µg/m³",
        "norma":  "NOM-172-SEMARNAT-2023 / NOM-025-SSA1-2021",
        "desc": (
            "Partículas con diámetro aerodinámico \u2264 10 µm, de origen natural "
            "(polvo mineral, suelo resuspendido, polen) y antrópico (tráfico, industria, "
            "agricultura). Penetran la nariz y la garganta y pueden alcanzar los bronquios. "
            "Umbrales NOM-172-SEMARNAT-2023 (prom. móvil pond. 12 h): "
            "Mala = 132 µg/m³ | Muy Mala = 213 µg/m³."
        ),
    },
    "PM25": {
        "nombre": "Partículas suspendidas finas (PM2.5)",
        "unidad": "µg/m³",
        "norma":  "NOM-172-SEMARNAT-2023 / NOM-025-SSA1-2021",
        "desc": (
            "Fracción con diámetro \u2264 2.5 µm, principalmente de combustión "
            "(vehículos, generación eléctrica, quemas). Alcanzan los alvéolos "
            "pulmonares y pueden pasar al torrente sanguíneo, con efectos "
            "cardiovasculares documentados a largo plazo. "
            "Umbrales NOM-172-SEMARNAT-2023 (prom. móvil pond. 12 h): "
            "Mala = 79 µg/m³ | Muy Mala = 130 µg/m³."
        ),
    },
    "SO2": {
        "nombre": "Dióxido de azufre (SO\u2082)",
        "unidad": "ppbv",
        "norma":  "NOM-172-SEMARNAT-2023 / NOM-022-SSA1-2019",
        "desc": (
            "Gas incoloro de olor acre emitido por la combustión de combustibles fósiles "
            "con alto contenido de azufre (refinación de petróleo, termoeléctricas). "
            "Produce irritación de las vías respiratorias superiores y contribuye a la "
            "formación de lluvia ácida. "
            "Umbrales NOM-172-SEMARNAT-2023 (promedio 1 h): "
            "Mala = 185 ppb | Muy Mala = 304 ppb."
        ),
    },
}

CIUDADES_DOMINIO = [
    "CDMX", "Cuernavaca", "Pachuca", "Puebla",
    "SJdelRio", "Tlaxcala", "Toluca", "Tula",
]
_CIUDADES_NORM = {c.lower(): c for c in CIUDADES_DOMINIO}
_CIUDADES_NORM.update({
    "sjdelrio": "SJdelRio", "san juan del rio": "SJdelRio",
    "cdmx": "CDMX", "ciudad de mexico": "CDMX",
})

HORIZONTES   = {"mod_dia1": "+24 h", "mod_dia2": "+48 h", "mod_dia3": "+72 h"}
HOR_LABELS   = ["+24 h", "+48 h", "+72 h"]
STAT_NAMES   = ["POD", "FAR", "CSI", "TSS", "PC", "BIAS"]
MIN_DIAS     = 5

# ── Control de calidad de datos observados (datos crudos SINAICA) ─────────────
#
# 1. Límites físicos mínimos y máximos por contaminante.
#    El par (obs, mod) se descarta si cualquier valor queda fuera del rango.
#
LIMITES_VALIDOS = {
    #          (min_obs, max_obs,  min_mod, max_mod)
    "O3":   (0.0,  300.0,  0.0,  300.0),
    "PM10": (0.0, 1000.0,  0.0, 1000.0),
    "PM25": (0.0,  500.0,  0.0,  500.0),   # techo absoluto; IQR lo refina
    "SO2":  (0.0, 1000.0,  0.0, 1000.0),
}

# 2. Filtro IQR sobre observaciones — solo para los contaminantes listados.
#    Elimina pares donde obs < Q1 - k·IQR  o  obs > Q3 + k·IQR.
#    k=3.0 (conservador) | k=1.5 (boxplot estándar, más agresivo).
#
CONTAMINANTES_FILTRO_IQR = {"PM25"}
IQR_FACTOR = 3.0

NOMBRE_MES = {
    "01": "Enero",   "02": "Febrero",  "03": "Marzo",
    "04": "Abril",   "05": "Mayo",     "06": "Junio",
    "07": "Julio",   "08": "Agosto",   "09": "Septiembre",
    "10": "Octubre", "11": "Noviembre","12": "Diciembre",
}

# ──────── colores (RGB) ─────────────────────────────────────────────────────
C_AZUL_OSC  = RGBColor(0x1F, 0x49, 0x7D)
C_AZUL_MED  = RGBColor(0x2E, 0x74, 0xB5)
C_AZUL_CLAR = RGBColor(0xBD, 0xD7, 0xEE)
C_TEAL      = RGBColor(0x21, 0x58, 0x68)
C_BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRIS_TXT  = RGBColor(0x44, 0x44, 0x44)

BG_HDR_OSC  = "1F497D"   # encabezado principal
BG_HDR_CLAR = "BDD7EE"   # subencabezado
BG_CIUDAD   = "EBF3FB"   # primera columna
BG_VERDE    = "C6EFCE"
BG_AMBAR    = "FFEB9C"
BG_ROJO     = "FFC7CE"
BG_GRIS     = "E8E8E8"
BG_BLANCO   = "FFFFFF"


# ──────────────────────────────────────────────────────────────────────────────
# UTILIDADES python-docx
# ──────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.replace("#", ""))
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_cell_borders(cell, color="AAAAAA", sz=4):
    """Aplica bordes finos a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_col_widths(table, widths_cm: List[float]):
    """Fija el ancho de cada columna en cm."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Eliminar tblW anterior
    old = tblPr.find(qn("w:tblW"))
    if old is not None:
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    total = int(sum(w * 567 for w in widths_cm))   # 567 twips ≈ 1 cm
    tblW.set(qn("w:w"),    str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # Filas: asignar ancho a cada celda según su índice lógico
    for row in table.rows:
        cells_seen = []
        for cell in row.cells:
            # Evitar celdas repetidas por merge horizontal
            if cell not in cells_seen:
                cells_seen.append(cell)
        for idx, cell in enumerate(cells_seen):
            if idx >= len(widths_cm):
                break
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            old_w = tcPr.find(qn("w:tcW"))
            if old_w is not None:
                tcPr.remove(old_w)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(int(widths_cm[idx] * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.insert(0, tcW)


def _write_cell(cell, text: str, bold=False, font_size=8,
                color: RGBColor = C_GRIS_TXT,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bg: str = BG_BLANCO,
                v_align=WD_ALIGN_VERTICAL.CENTER):
    """Escribe texto en una celda con formato completo."""
    cell.vertical_alignment = v_align
    _set_cell_bg(cell, bg)
    _set_cell_borders(cell)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color


def _semaforo_pod(v) -> str:
    if v is None: return BG_GRIS
    return BG_VERDE if v >= 0.7 else BG_AMBAR if v >= 0.4 else BG_ROJO

def _semaforo_far(v) -> str:
    if v is None: return BG_GRIS
    return BG_VERDE if v <= 0.3 else BG_AMBAR if v <= 0.5 else BG_ROJO

def _semaforo_csi(v) -> str:
    if v is None: return BG_GRIS
    return BG_VERDE if v >= 0.4 else BG_AMBAR if v >= 0.2 else BG_ROJO

def _semaforo_tss(v) -> str:
    if v is None: return BG_GRIS
    return BG_VERDE if v >= 0.4 else BG_AMBAR if v >= 0.1 else BG_ROJO

def _semaforo_bias(v) -> str:
    if v is None: return BG_GRIS
    d = abs(v - 1.0)
    return BG_VERDE if d <= 0.3 else BG_AMBAR if d <= 0.6 else BG_ROJO

SEMAFOROS = {
    "POD":  _semaforo_pod,
    "FAR":  _semaforo_far,
    "CSI":  _semaforo_csi,
    "TSS":  _semaforo_tss,
    "PC":   lambda v: BG_BLANCO,
    "BIAS": _semaforo_bias,
}

def _fmt(v) -> str:
    return "N/D" if v is None else f"{v:.3f}"


# ──────────────────────────────────────────────────────────────────────────────
# CONTROL DE CALIDAD DE DATOS
# ──────────────────────────────────────────────────────────────────────────────

def limpiar_serie(
    obs: np.ndarray,
    mod: np.ndarray,
    contaminante: str,
) -> tuple:
    """
    Aplica control de calidad a un par (obs, mod) antes de calcular
    la tabla de contingencia.

    Pasos (en orden):
        1. Eliminar NaN / Inf en obs o mod.
        2. Límites físicos por contaminante (LIMITES_VALIDOS):
           descarta el par si obs o mod quedan fuera del rango válido.
           Captura valores negativos e implausiblemente altos.
        3. Filtro IQR sobre obs (solo contaminantes en CONTAMINANTES_FILTRO_IQR):
           descarta el par si obs < Q1 - k·IQR  o  obs > Q3 + k·IQR.

    Retorna
    -------
    obs_limpio, mod_limpio : arrays filtrados
    resumen : dict con n_original, n_nan, n_limites, n_iqr, n_final
    """
    n_original = len(obs)

    # Paso 1: NaN / Inf
    mask = np.isfinite(obs) & np.isfinite(mod)
    obs, mod = obs[mask], mod[mask]
    n_nan = n_original - len(obs)

    # Paso 2: límites físicos
    lim = LIMITES_VALIDOS.get(contaminante)
    n_limites = 0
    if lim is not None:
        min_obs, max_obs, min_mod, max_mod = lim
        mask_lim = (
            (obs >= min_obs) & (obs <= max_obs) &
            (mod >= min_mod) & (mod <= max_mod)
        )
        n_antes = len(obs)
        obs, mod = obs[mask_lim], mod[mask_lim]
        n_limites = n_antes - len(obs)

    # Paso 3: filtro IQR sobre obs (solo PM2.5 por defecto)
    n_iqr = 0
    if contaminante in CONTAMINANTES_FILTRO_IQR and len(obs) >= 4:
        q1, q3 = np.percentile(obs, [25, 75])
        iqr = q3 - q1
        if iqr > 0:
            lim_inf = q1 - IQR_FACTOR * iqr
            lim_sup = q3 + IQR_FACTOR * iqr
            mask_iqr = (obs >= lim_inf) & (obs <= lim_sup)
            n_antes = len(obs)
            obs, mod = obs[mask_iqr], mod[mask_iqr]
            n_iqr = n_antes - len(obs)

    resumen = {
        "n_original": n_original,
        "n_nan":      n_nan,
        "n_limites":  n_limites,
        "n_iqr":      n_iqr,
        "n_final":    len(obs),
    }
    return obs, mod, resumen


# ──────────────────────────────────────────────────────────────────────────────
# CÁLCULO DE ESTADÍSTICOS DICOTÓMICOS
# ──────────────────────────────────────────────────────────────────────────────

def contingencia(
    obs: np.ndarray,
    mod: np.ndarray,
    umbral: float,
    contaminante: str = "",
) -> Optional[Dict]:
    """
    Aplica control de calidad y calcula la tabla de contingencia 2×2.

    Parámetros
    ----------
    obs           : array de observaciones (valores crudos SINAICA)
    mod           : array del modelo (mismo horizonte)
    umbral        : valor normativo de referencia (evento = obs/mod >= umbral)
    contaminante  : clave del contaminante para seleccionar límites y filtro IQR

    Retorna dict con H, M, F, C, N, POD, FAR, CSI, TSS, PC, BIAS
    y columnas QC: n_orig, n_nan, n_lim, n_iqr.
    Retorna None si N < MIN_DIAS tras el filtro.
    """
    # ── Control de calidad ────────────────────────────────────────────────────
    obs, mod, qc = limpiar_serie(
        np.asarray(obs, float), np.asarray(mod, float), contaminante
    )

    N = qc["n_final"]
    if N < MIN_DIAS:
        return None

    # ── Tabla de contingencia ─────────────────────────────────────────────────
    obs_ev = obs >= umbral
    mod_ev = mod >= umbral
    H = int(np.sum( obs_ev &  mod_ev))
    M = int(np.sum( obs_ev & ~mod_ev))
    F = int(np.sum(~obs_ev &  mod_ev))
    C = int(np.sum(~obs_ev & ~mod_ev))

    def safe(num, den):
        return num / den if den > 0 else np.nan

    POD  = safe(H, H + M)
    FAR  = safe(F, H + F)
    CSI  = safe(H, H + M + F)
    POFD = safe(F, F + C)
    TSS  = POD - POFD
    PC   = safe(H + C, N)
    BIAS = safe(H + F, H + M)

    def r(x): return round(float(x), 3) if np.isfinite(x) else None

    return {
        # Contingencia
        "N": N, "H": H, "M": M, "F": F, "C": C,
        # Métricas
        "POD": r(POD), "FAR": r(FAR), "CSI": r(CSI),
        "TSS": r(TSS), "PC":  r(PC),  "BIAS": r(BIAS),
        # Control de calidad
        "n_orig": qc["n_original"],
        "n_nan":  qc["n_nan"],
        "n_lim":  qc["n_limites"],
        "n_iqr":  qc["n_iqr"],
    }


# ──────────────────────────────────────────────────────────────────────────────
# LECTURA DE CSV Y CÁLCULO
# ──────────────────────────────────────────────────────────────────────────────

def normalizar_ciudad(nombre: str) -> Optional[str]:
    return _CIUDADES_NORM.get(nombre.strip().lower())

def parsear_lista_ciudades(valor) -> Optional[List[str]]:
    if not valor:
        return None
    crudos = []
    items = [valor] if isinstance(valor, str) else valor
    for item in items:
        crudos.extend(t for t in re.split(r"[,\s]+", item.strip()) if t)
    canonicas, invalidas = [], []
    for c in crudos:
        cn = normalizar_ciudad(c)
        if cn is None: invalidas.append(c)
        elif cn not in canonicas: canonicas.append(cn)
    if invalidas:
        sys.exit(f"[ERROR] Ciudad(es) no reconocida(s): {invalidas}\n"
                 f"        Válidas: {CIUDADES_DOMINIO}")
    return canonicas

def parsear_nombre(ruta: str):
    m = re.match(r"^eval_([A-Za-z0-9]+)_(.+)_\d{4}-\d{2}-\d{2}$", Path(ruta).stem)
    if not m: return None
    cont   = m.group(1).upper()
    ciudad = normalizar_ciudad(m.group(2)) or m.group(2)
    return cont, ciudad

def leer_csvs(directorio: str, ciudades_filtro: Optional[List[str]], mes_filtro: Optional[str]) -> dict:
    archivos = sorted(glob.glob(os.path.join(directorio, "eval_*.csv")))
    if not archivos:
        sys.exit(f"[ERROR] No se encontraron eval_*.csv en '{directorio}'.")
    datos = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    n_ok  = 0
    for ruta in archivos:
        meta = parsear_nombre(ruta)
        if meta is None: continue
        cont, ciudad = meta
        if cont not in UMBRALES: continue
        if ciudades_filtro and ciudad not in ciudades_filtro: continue
        try:
            df = pd.read_csv(ruta, parse_dates=["Fecha"])
        except Exception: continue
        if not {"Fecha","max_obs","mod_dia1","mod_dia2","mod_dia3"}.issubset(df.columns): continue
        mes = df["Fecha"].dt.strftime("%Y-%m").iloc[0]
        if mes_filtro and mes != mes_filtro: continue
        datos[mes][ciudad][cont].append(df)
        n_ok += 1
    print(f"[INFO] Archivos leídos: {n_ok}")
    if n_ok == 0:
        sys.exit("[ERROR] Ningún archivo coincidió con los filtros.")
    return datos

def calcular_resultados(datos: dict) -> dict:
    """
    Calcula los estadísticos dicotómicos para cada mes/ciudad/cont/horizonte.
    Aplica control de calidad (limpiar_serie) antes de cada cálculo y
    reporta en el log los pares descartados con etiqueta [QC].
    """
    resultados  = defaultdict(lambda: defaultdict(lambda: defaultdict(dict)))
    # Acumuladores globales para el resumen final
    qc_resumen: Dict[str, Dict[str, int]] = {}

    for mes, ciudades in sorted(datos.items()):
        qc_mes = {"n_nan": 0, "n_lim": 0, "n_iqr": 0}

        for ciudad, conts in sorted(ciudades.items()):
            for cont, lista_df in sorted(conts.items()):
                umbral = UMBRALES[cont]
                df_mes = pd.concat(lista_df, ignore_index=True).sort_values("Fecha")
                obs    = np.asarray(df_mes["max_obs"].values, float)

                for col, hor in HORIZONTES.items():
                    mod = np.asarray(df_mes[col].values, float)
                    # Pasar contaminante para aplicar QC específico
                    st  = contingencia(obs, mod, umbral, contaminante=cont)

                    resultados[mes][cont][ciudad][hor] = st

                    if st is None:
                        continue

                    # Reportar descarte si hubo filtrado
                    total_desc = st["n_nan"] + st["n_lim"] + st["n_iqr"]
                    if total_desc > 0:
                        desc = []
                        if st["n_nan"] > 0:
                            desc.append(f"NaN/Inf={st['n_nan']}")
                        if st["n_lim"] > 0:
                            desc.append(f"límites físicos={st['n_lim']}")
                        if st["n_iqr"] > 0:
                            desc.append(f"outliers IQR={st['n_iqr']}")
                        print(
                            f"[QC]  {mes} | {ciudad} {cont} {hor}: "
                            f"{total_desc} pares descartados "
                            f"({', '.join(desc)}) → quedan {st['N']}/{st['n_orig']}"
                        )
                        qc_mes["n_nan"] += st["n_nan"]
                        qc_mes["n_lim"] += st["n_lim"]
                        qc_mes["n_iqr"] += st["n_iqr"]

        total_mes = sum(qc_mes.values())
        if total_mes > 0:
            print(
                f"[QC]  {mes}: resumen del mes — "
                f"NaN/Inf={qc_mes['n_nan']}, "
                f"límites físicos={qc_mes['n_lim']}, "
                f"outliers IQR={qc_mes['n_iqr']}"
            )
        qc_resumen[mes] = qc_mes

    return resultados

def exportar_csv(resultados: dict, ruta: str):
    """
    Exporta todos los estadísticos dicotómicos a CSV de auditoría.
    Incluye columnas de control de calidad: n_orig, n_nan, n_lim, n_iqr.
    """
    filas = []
    for mes, conts in sorted(resultados.items()):
        for cont, ciudades in sorted(conts.items()):
            for ciudad, hors in sorted(ciudades.items()):
                for hor, st in sorted(hors.items()):
                    if st is None:
                        continue
                    filas.append({
                        "mes":          mes,
                        "contaminante": cont,
                        "ciudad":       ciudad,
                        "horizonte":    hor,
                        # Control de calidad
                        "N":      st["N"],
                        "n_orig": st["n_orig"],
                        "n_nan":  st["n_nan"],
                        "n_lim":  st["n_lim"],
                        "n_iqr":  st["n_iqr"],
                        # Tabla de contingencia
                        "H": st["H"], "M": st["M"],
                        "F": st["F"], "C": st["C"],
                        # Métricas dicotómicas
                        "POD":  st["POD"],  "FAR":  st["FAR"],
                        "CSI":  st["CSI"],  "TSS":  st["TSS"],
                        "PC":   st["PC"],   "BIAS": st["BIAS"],
                    })
    if filas:
        pd.DataFrame(filas).to_csv(ruta, index=False)
        print(f"[OK]  CSV de auditoría: {ruta}  ({len(filas)} filas)")


# ──────────────────────────────────────────────────────────────────────────────
# CONSTRUCCIÓN DEL DOCUMENTO WORD
# ──────────────────────────────────────────────────────────────────────────────

def _page_landscape(doc: Document):
    """Configura orientación horizontal A4 para todas las secciones."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    # Forzar orientación landscape en el XML
    pgSz = section._sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        pgSz.set(qn("w:orient"), "landscape")


def _add_heading(doc, text: str, level: int, color: RGBColor = C_AZUL_OSC,
                 size: int = 14, space_before: int = 12, space_after: int = 6):
    para = doc.add_paragraph()
    para.style = doc.styles[f"Heading {level}"]
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    run = para.add_run(text)
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return para


def _add_body(doc, text: str, size: int = 9, italic: bool = False,
              space_before: int = 3, space_after: int = 3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.italics   = italic
    return para


def _tabla_umbrales(doc, cont_sections: list):
    """Tabla resumen de contaminantes y umbrales normativos."""
    headers  = ["Contaminante", "Umbral", "Unidad", "Norma", "Descripción breve"]
    col_w    = [4.5, 1.5, 1.5, 3.5, 15.5]
    table    = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Encabezado
    for i, (h, w) in enumerate(zip(headers, col_w)):
        _write_cell(table.rows[0].cells[i], h, bold=True, font_size=8,
                    color=C_BLANCO, bg=BG_HDR_OSC)
    _set_col_widths(table, col_w)
    # Filas de datos
    for cs in cont_sections:
        row = table.add_row()
        vals = [cs["nombre"], str(cs["umbral"]), cs["unidad"], cs["norma"], cs["desc_corta"]]
        bgs  = [BG_CIUDAD, BG_AZUL_CLAR := "D9E8F5", BG_BLANCO, BG_BLANCO, BG_BLANCO]
        for i, (v, bg) in enumerate(zip(vals, bgs)):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER
            _write_cell(row.cells[i], v, font_size=8, bg=bg, align=align)
    _set_col_widths(table, col_w)
    return table


def _tabla_contingencia_conceptual(doc):
    """Tabla 2×2 conceptual H/M/F/C."""
    col_w = [4.0, 4.5, 4.5]
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    datos = [
        ["",               "Obs. EVENTO",    "Obs. NO EVENTO"],
        ["Mod. EVENTO",    "H — Acierto",    "F — Falsa Alarma"],
        ["Mod. NO EVENTO", "M — Fallo",      "C — Rechazo Correcto"],
    ]
    bgs = [
        [BG_HDR_OSC, BG_HDR_OSC, BG_HDR_OSC],
        [BG_CIUDAD,  BG_VERDE,   BG_ROJO],
        [BG_CIUDAD,  BG_ROJO,    BG_VERDE],
    ]
    colores_txt = [
        [C_BLANCO,   C_BLANCO,   C_BLANCO],
        [C_TEAL,     RGBColor(0x1a,0x7a,0x3c), RGBColor(0x9c,0x1b,0x1b)],
        [C_TEAL,     RGBColor(0x9c,0x1b,0x1b), RGBColor(0x1a,0x7a,0x3c)],
    ]
    for r, (fila, bg_fila, col_fila) in enumerate(zip(datos, bgs, colores_txt)):
        for c, (txt, bg, col) in enumerate(zip(fila, bg_fila, col_fila)):
            _write_cell(table.rows[r].cells[c], txt, bold=True, font_size=8,
                        color=col, bg=bg)
    _set_col_widths(table, col_w)
    return table


def _tabla_leyenda(doc):
    """Tabla de semáforo de desempeño."""
    col_w = [1.8, 3.0, 3.0, 3.0]
    tabla_data = [
        ("POD",  "≥ 0.700",       "0.400 – 0.699", "< 0.400"),
        ("FAR",  "≤ 0.300",       "0.301 – 0.500", "> 0.500"),
        ("CSI",  "≥ 0.400",       "0.200 – 0.399", "< 0.200"),
        ("TSS",  "≥ 0.400",       "0.100 – 0.399", "< 0.100"),
        ("BIAS", "|BIAS-1| ≤ 0.3","0.3 – 0.6",     "> 0.6"),
    ]
    table = doc.add_table(rows=1 + len(tabla_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Métrica", "🟢  Bueno", "🟡  Aceptable", "🔴  Deficiente"]):
        _write_cell(table.rows[0].cells[i], h, bold=True, font_size=8,
                    color=C_BLANCO, bg=BG_HDR_OSC)
    for ri, (stat, verde, ambar, rojo) in enumerate(tabla_data, start=1):
        row = table.rows[ri]
        _write_cell(row.cells[0], stat,  bold=True, font_size=8, bg=BG_CIUDAD, color=C_TEAL)
        _write_cell(row.cells[1], verde, font_size=8, bg=BG_VERDE,
                    color=RGBColor(0x1a,0x7a,0x3c))
        _write_cell(row.cells[2], ambar, font_size=8, bg=BG_AMBAR,
                    color=RGBColor(0x7d,0x5a,0x00))
        _write_cell(row.cells[3], rojo,  font_size=8, bg=BG_ROJO,
                    color=RGBColor(0x9c,0x1b,0x1b))
    _set_col_widths(table, col_w)
    return table


def _tabla_estadisticos(doc, cont_data: dict, ciudades: list):
    """
    Tabla principal de estadísticos (POD FAR CSI TSS PC BIAS) × 3 horizontes.
    Layout landscape A4 (área útil ≈ 26.7 cm):
      Ciudad(3.5) + N(1.2) + 6stats×3hor×(22/18=1.222cm) = 3.5+1.2+22 = 26.7
    """
    N_STATS = len(STAT_NAMES)
    N_HOR   = len(HOR_LABELS)
    # Anchuras en cm
    W_CIUDAD = 3.2
    W_N      = 1.0
    W_STAT   = (26.7 - W_CIUDAD - W_N) / (N_STATS * N_HOR)  # ≈1.21 cm

    n_cols = 2 + N_STATS * N_HOR
    table  = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ── Fila 1: encabezado principal ─────────────────────────────────────────
    row1 = table.add_row()
    _write_cell(row1.cells[0], "Ciudad",  bold=True, font_size=7.5,
                color=C_BLANCO, bg=BG_HDR_OSC)
    _write_cell(row1.cells[1], "N días",  bold=True, font_size=7.5,
                color=C_BLANCO, bg=BG_HDR_OSC)
    col = 2
    for hor in HOR_LABELS:
        # Combinar N_STATS celdas con merge
        start_cell = row1.cells[col]
        end_cell   = row1.cells[col + N_STATS - 1]
        merged     = start_cell.merge(end_cell)
        _write_cell(merged, hor, bold=True, font_size=8, color=C_BLANCO, bg=BG_HDR_OSC)
        col += N_STATS

    # ── Fila 2: nombres de estadísticos ──────────────────────────────────────
    row2 = table.add_row()
    _write_cell(row2.cells[0], "", font_size=7, bg=BG_HDR_CLAR)
    _write_cell(row2.cells[1], "", font_size=7, bg=BG_HDR_CLAR)
    col = 2
    for _ in range(N_HOR):
        for stat in STAT_NAMES:
            _write_cell(row2.cells[col], stat, bold=True, font_size=7,
                        color=C_AZUL_OSC, bg=BG_HDR_CLAR)
            col += 1

    # ── Filas de datos por ciudad ─────────────────────────────────────────────
    for ciudad in ciudades:
        hors_data = cont_data.get(ciudad, {})
        h24 = hors_data.get("+24 h")
        n_dias = str(h24["N"]) if h24 else "—"

        row = table.add_row()
        _write_cell(row.cells[0], ciudad, bold=True, font_size=8,
                    color=C_TEAL, bg=BG_CIUDAD,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
        _write_cell(row.cells[1], n_dias, bold=True, font_size=8,
                    color=C_TEAL, bg=BG_CIUDAD)
        col = 2
        for hor in HOR_LABELS:
            st = hors_data.get(hor)
            for stat in STAT_NAMES:
                val = st.get(stat) if st else None
                bg  = SEMAFOROS[stat](val)
                _write_cell(row.cells[col], _fmt(val), font_size=7.5, bg=bg)
                col += 1

    # Anchos finales
    _set_col_widths(table, [W_CIUDAD, W_N] + [W_STAT] * (N_STATS * N_HOR))
    return table


def _tabla_contingencia_hMFC(doc, cont_data: dict, ciudades: list):
    """
    Tabla de valores crudos H, M, F, C por ciudad × 3 horizontes.
    """
    W_CIUDAD = 3.2
    W_N      = 1.0
    W_HMFC   = (26.7 - W_CIUDAD - W_N) / (4 * 3)   # 4 celdas × 3 horizontes

    n_cols = 1 + 4 * 3
    table  = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado principal
    row1 = table.add_row()
    _write_cell(row1.cells[0], "Ciudad", bold=True, font_size=7.5,
                color=C_BLANCO, bg=BG_HDR_OSC)
    col = 1
    for hor in HOR_LABELS:
        merged = row1.cells[col].merge(row1.cells[col + 3])
        _write_cell(merged, hor, bold=True, font_size=8, color=C_BLANCO, bg=BG_HDR_OSC)
        col += 4

    # Subencabezado H M F C
    row2 = table.add_row()
    _write_cell(row2.cells[0], "", font_size=7, bg=BG_HDR_CLAR)
    lbl_bg = {"H": BG_VERDE, "M": BG_ROJO, "F": BG_ROJO, "C": BG_VERDE}
    col = 1
    for _ in range(3):
        for lbl in ["H", "M", "F", "C"]:
            _write_cell(row2.cells[col], lbl, bold=True, font_size=7,
                        color=C_AZUL_OSC, bg=BG_HDR_CLAR)
            col += 1

    # Filas de datos
    for ciudad in ciudades:
        hors_data = cont_data.get(ciudad, {})
        row = table.add_row()
        _write_cell(row.cells[0], ciudad, bold=True, font_size=8,
                    color=C_TEAL, bg=BG_CIUDAD, align=WD_ALIGN_PARAGRAPH.LEFT)
        col = 1
        for hor in HOR_LABELS:
            st = hors_data.get(hor)
            for lbl in ["H", "M", "F", "C"]:
                val  = str(st[lbl]) if st else "—"
                bg   = lbl_bg[lbl] if st and st[lbl] > 0 else BG_BLANCO
                _write_cell(row.cells[col], val, font_size=7.5, bg=bg)
                col += 1

    col_widths = [W_CIUDAD] + [W_HMFC] * 12
    _set_col_widths(table, col_widths)
    return table


def _add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


def generar_docx(resultados: dict, ruta_salida: str, cat_key: str = "mala"):
    """Construye y guarda el documento Word completo."""
    doc = Document()
    _page_landscape(doc)

    # Etiquetas de la categoría para el documento
    cat_label  = "Mala (Alto / naranja)"      if cat_key == "mala" else "Muy Mala (Muy Alto / rojo)"
    cat_color  = RGBColor(0xFF, 0x7E, 0x00)   if cat_key == "mala" else RGBColor(0xFF, 0x00, 0x00)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("EVALUACIÓN DICOTÓMICA MENSUAL DEL PRONÓSTICO")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = C_AZUL_OSC

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("WRF-Chem vs SINAICA/INECC")
    r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = C_AZUL_MED

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Centro de México — Ocho zonas metropolitanas")
    r3.font.size = Pt(12); r3.font.color.rgb = C_TEAL

    meses_txt = "  ·  ".join(
        NOMBRE_MES.get(m.split("-")[1], m) + " " + m.split("-")[0]
        for m in sorted(resultados.keys())
    )
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(meses_txt)
    r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(0x44,0x44,0x44)

    # Categoría de evaluación en portada
    p_cat = doc.add_paragraph()
    p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cat.paragraph_format.space_before = Pt(8)
    r_cat_lbl = p_cat.add_run("Categoría de evaluación (NOM-172-SEMARNAT-2023): ")
    r_cat_lbl.font.size = Pt(11); r_cat_lbl.font.color.rgb = RGBColor(0x44,0x44,0x44)
    r_cat_val = p_cat.add_run(cat_label)
    r_cat_val.bold = True; r_cat_val.font.size = Pt(11); r_cat_val.font.color.rgb = cat_color

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(
        "Generado automáticamente por informe_dicotomico.py v1.3.0 — ddsinaica / ICAyCC, UNAM"
    )
    r5.font.size = Pt(8); r5.font.color.rgb = RGBColor(0x88,0x88,0x88)

    # Línea separadora
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(12)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBdr.append(bot)
    pPr.append(pBdr)

    # ── SECCIÓN 1: INTRODUCCIÓN ───────────────────────────────────────────────
    _add_heading(doc, "1. Introducción y metodología", 1, C_AZUL_OSC, 13, 14, 6)
    _add_body(doc,
        "El presente informe resume el desempeño del sistema de pronóstico de calidad "
        "del aire WRF-Chem en la detección de episodios de contaminación para cuatro "
        "contaminantes regulados en ocho zonas metropolitanas del centro de México. "
        "Los umbrales de clasificación corresponden a la NOM-172-SEMARNAT-2023 "
        "(DOF 25/01/2024), Índice AIRE Y SALUD. La evaluación dicotómica compara si "
        "el máximo diario observado y el modelado superan o no el umbral de la "
        f"categoría seleccionada: {cat_label}.",
        size=9, space_before=4, space_after=8)

    _add_heading(doc, "1.1  Umbrales NOM-172-SEMARNAT-2023", 2, C_AZUL_MED, 11, 8, 4)

    # Tabla comparativa de ambas categorías
    BG_MALA     = "FFE0B2"   # naranja claro
    BG_MUY_MALA = "FFCDD2"   # rojo claro
    BG_ACTIVA   = "FFE0B2" if cat_key == "mala" else "FFCDD2"

    col_w_nom = [4.5, 1.5, 3.0, 3.0, 1.5]
    tbl_nom = doc.add_table(rows=1, cols=5)
    tbl_nom.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdrs = ["Contaminante", "Unidad",
            "🟠 Mala (Alto)",
            "🔴 Muy Mala (Muy Alto)",
            "Usada aquí"]
    for i, h in enumerate(hdrs):
        bg = BG_ACTIVA if i in ((2,) if cat_key == "mala" else (3,)) else BG_HDR_OSC
        if i < 2 or i == 4:
            bg = BG_HDR_OSC
        _write_cell(tbl_nom.rows[0].cells[i], h, bold=True, font_size=8,
                    color=C_BLANCO if bg == BG_HDR_OSC else RGBColor(0x33,0x33,0x33),
                    bg=bg)
    _set_col_widths(tbl_nom, col_w_nom)

    for cont in ["O3", "PM10", "PM25", "SO2"]:
        if not any(cont in resultados[mes] for mes in resultados):
            continue
        meta = META_CONT[cont]
        val_mala     = CATEGORIAS_NOM172["mala"][cont]
        val_muy_mala = CATEGORIAS_NOM172["muy_mala"][cont]
        activa_txt   = f"{'mala' if cat_key == 'mala' else 'muy_mala'} ({UMBRALES[cont]} {meta['unidad']})"

        row = tbl_nom.add_row()
        _write_cell(row.cells[0], meta["nombre"],         font_size=8, bg=BG_CIUDAD,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
        _write_cell(row.cells[1], meta["unidad"],         font_size=8)
        _write_cell(row.cells[2], f"{val_mala}",          font_size=8, bg=BG_MALA)
        _write_cell(row.cells[3], f"{val_muy_mala}",      font_size=8, bg=BG_MUY_MALA)
        _write_cell(row.cells[4], activa_txt,             font_size=8, bold=True,
                    bg=BG_ACTIVA,
                    color=RGBColor(0x7F,0x33,0x00) if cat_key=="mala" else RGBColor(0x7F,0x00,0x00))
    _set_col_widths(tbl_nom, col_w_nom)
    doc.add_paragraph()

    _add_heading(doc, "1.2  Tabla de contingencia 2×2", 2, C_AZUL_MED, 11, 8, 4)
    _add_body(doc,
        "Cada día se clasifica como EVENTO (valor ≥ umbral) o NO EVENTO. "
        "La intersección entre observación y modelo produce cuatro categorías:",
        size=9, space_before=3, space_after=5)
    _tabla_contingencia_conceptual(doc)
    doc.add_paragraph()

    _add_heading(doc, "1.3  Definición de métricas", 2, C_AZUL_MED, 11, 8, 4)
    defs = [
        ("POD",  "= H / (H + M)",
         "Probabilidad de detección. Fracción de eventos observados correctamente detectados. Ideal = 1."),
        ("FAR",  "= F / (H + F)",
         "Tasa de falsas alarmas. Fracción de eventos pronosticados que no ocurrieron. Ideal = 0."),
        ("CSI",  "= H / (H + M + F)",
         "Índice de éxito crítico. Integra fallos y falsas alarmas; ignora rechazos correctos. Ideal = 1."),
        ("TSS",  "= H/(H+M) − F/(F+C)",
         "Pierce Skill Score (Hanssen-Kuipers). Diferencia entre tasa de detección y tasa de falsa detección. Rango [−1, 1]; ideal = 1."),
        ("PC",   "= (H + C) / N",
         "Porcentaje correcto. Fracción de días correctamente clasificados. Ideal = 1. Puede ser engañoso cuando los eventos son raros."),
        ("BIAS", "= (H + F) / (H + M)",
         "Sesgo de frecuencia. Cociente entre eventos pronosticados y observados. Ideal = 1; >1 sobreestima; <1 subestima."),
    ]
    for nombre, formula, desc in defs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(1)
        r_n = p.add_run(f"{nombre} ")
        r_n.bold = True; r_n.font.size = Pt(9); r_n.font.color.rgb = C_AZUL_OSC
        r_f = p.add_run(formula + "   ")
        r_f.italic = True; r_f.font.size = Pt(9); r_f.font.color.rgb = C_AZUL_MED
        r_d = p.add_run(desc)
        r_d.font.size = Pt(8.5); r_d.font.color.rgb = C_GRIS_TXT

    doc.add_paragraph()
    _add_heading(doc, "1.4  Leyenda de semáforo de desempeño", 2, C_AZUL_MED, 11, 8, 4)
    _tabla_leyenda(doc)

    # ── SECCIONES POR MES ────────────────────────────────────────────────────
    sec_num = 2
    for mes in sorted(resultados.keys()):
        anio, mm = mes.split("-")
        nombre_mes = NOMBRE_MES.get(mm, mm)

        _add_page_break(doc)
        _add_heading(doc,
            f"{sec_num}. Resultados — {nombre_mes} {anio}",
            1, C_AZUL_OSC, 13, 14, 6)

        # Obtener ciudades presentes en este mes (unión de todos los contaminantes)
        ciudades_mes = sorted({
            ciudad
            for cont in resultados[mes].values()
            for ciudad in cont.keys()
        })

        sub = 1
        for cont in ["O3", "PM10", "PM25", "SO2"]:
            if cont not in resultados[mes]:
                continue
            meta   = META_CONT[cont]
            umbral = UMBRALES[cont]

            _add_heading(doc,
                f"{sec_num}.{sub}  {meta['nombre']}  "
                f"(umbral: {umbral} {meta['unidad']}, {meta['norma']})",
                2, C_AZUL_MED, 11, 10, 4)
            _add_body(doc, meta["desc"], size=8.5, space_before=3, space_after=8)

            # Ciudades con datos para este contaminante
            ciudades_cont = sorted(resultados[mes][cont].keys())

            _add_heading(doc, "Estadísticos de verificación dicotómica", 3,
                         C_TEAL, 10, 6, 3)
            _tabla_estadisticos(doc, resultados[mes][cont], ciudades_cont)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

            _add_heading(doc, "Tabla de contingencia (H, M, F, C)", 3,
                         C_TEAL, 10, 6, 3)
            _tabla_contingencia_hMFC(doc, resultados[mes][cont], ciudades_cont)
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

            sub += 1
        sec_num += 1

    # ── NOTAS FINALES ─────────────────────────────────────────────────────────
    _add_page_break(doc)
    _add_heading(doc, "Notas técnicas", 1, C_AZUL_OSC, 13, 14, 6)
    notas = [
        f"N/D indica que no se alcanzó el mínimo de {MIN_DIAS} días con datos válidos "
        "tras el control de calidad para calcular estadísticos confiables.",
        "Control de calidad sobre datos crudos SINAICA: (1) se eliminan pares con "
        "valores negativos o fuera de los límites físicos del contaminante; "
        "(2) sobre PM2.5 se aplica un filtro IQR (k·IQR sobre obs) que descarta "
        "observaciones extraordinarias no representativas de episodios reales. "
        "Los valores descartados se reportan en el log [QC] y en el CSV de auditoría "
        "(columnas n_orig, n_nan, n_lim, n_iqr).",
        "El horizonte +72 h cubre sólo 18 h de la ventana local (índices 54–71 del wrfout) "
        "por la alineación UTC−6 con el inicio del run de WRF-Chem.",
        "BIAS = 1 indica que el modelo pronosticó el mismo número de eventos que los observados "
        "(sin implicar coincidencia día a día). BIAS > 1: sobreestima eventos; BIAS < 1: los subestima.",
        "Las observaciones provienen de la red SINAICA/INECC; los valores del modelo corresponden "
        "al máximo diario extraído del dominio WRF-Chem (máximo espacial sobre la región de la ciudad).",
        "Código fuente: https://github.com/JoseAgustin/ddsinaica — informe_dicotomico.py v1.2.0",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(nota)
        r.font.size = Pt(8.5)

    doc.save(ruta_salida)
    print(f"[OK]  Documento guardado: {ruta_salida}")


# ──────────────────────────────────────────────────────────────────────────────
# PUNTO DE ENTRADA
# ──────────────────────────────────────────────────────────────────────────────

def parse_args():
    p = argparse.ArgumentParser(
        prog="informe_dicotomico.py",
        description=(
            "Genera un informe Word (.docx) con estadísticos dicotómicos "
            "mensuales (POD, FAR, CSI, TSS, PC, BIAS) para la evaluación "
            "WRF-Chem vs SINAICA/INECC. Sin dependencia de Node.js."
        ),
    )
    p.add_argument("--entrada",  "-i", default="combinado/ajustados",
                   help="Directorio con CSV eval_*.csv (default: combinado/ajustados)")
    p.add_argument("--salida",   "-o", default="informe_dicotomico.docx",
                   help="Ruta del .docx de salida")
    p.add_argument("--mes",      "-m", default=None,
                   help="Procesar solo este mes (YYYY-MM)")
    p.add_argument("--ciudades", "-c", nargs="+", default=None,
                   help=f"Ciudades a incluir. Catálogo: {', '.join(CIUDADES_DOMINIO)}")
    p.add_argument(
        "--categoria", "-k",
        default=CATEGORIA_DEFAULT,
        metavar="NIVEL",
        help=(
            "Categoría de calidad del aire para el cálculo dicotómico "
            "(evento = valor >= umbral de la categoría). "
            "Opciones: 'mala' (riesgo Alto / naranja, default) | "
            "'muy_mala' (riesgo Muy Alto / rojo). "
            "Fuente: NOM-172-SEMARNAT-2023."
        ),
    )
    p.add_argument("--csv-auditoria", default=None,
                   help="Ruta opcional para exportar CSV con todos los estadísticos")
    p.add_argument(
        "--umbral-pm25",
        type=float,
        default=LIMITES_VALIDOS["PM25"][1],
        metavar="UG_M3",
        help=(
            "Techo absoluto para observaciones de PM2.5 en µg/m³. "
            "Pares con max_obs > este valor se descartan. "
            f"Default: {LIMITES_VALIDOS['PM25'][1]} µg/m³."
        ),
    )
    p.add_argument(
        "--iqr-factor",
        type=float,
        default=IQR_FACTOR,
        metavar="K",
        help=(
            "Factor multiplicador del IQR para detección de outliers en PM2.5. "
            "Elimina obs fuera de [Q1 - k·IQR, Q3 + k·IQR]. "
            "k=3.0 (default, conservador) | k=1.5 (boxplot estándar, agresivo)."
        ),
    )
    return p.parse_args()


def main():
    args            = parse_args()
    ciudades_filtro = parsear_lista_ciudades(args.ciudades)

    # ── Resolver categoría ────────────────────────────────────────────────────
    cat_key = ALIAS_CATEGORIA.get(args.categoria.lower().replace(" ", "_"))
    if cat_key is None:
        sys.exit(
            f"[ERROR] Categoría no reconocida: '{args.categoria}'\n"
            f"        Opciones válidas: mala, muy_mala"
        )

    # Actualizar UMBRALES global con la categoría seleccionada
    global UMBRALES, IQR_FACTOR, LIMITES_VALIDOS
    UMBRALES = dict(CATEGORIAS_NOM172[cat_key])

    # ── Aplicar parámetros QC a las variables globales ────────────────────────
    IQR_FACTOR = args.iqr_factor
    lim_pm25   = list(LIMITES_VALIDOS["PM25"])
    lim_pm25[1] = args.umbral_pm25
    LIMITES_VALIDOS["PM25"] = tuple(lim_pm25)

    # ── Nombre de salida con sufijo de categoría si no fue especificado ───────
    salida = args.salida
    if salida == "informe_dicotomico.docx":
        salida = f"informe_dicotomico_{cat_key}.docx"

    print("=" * 62)
    print("  Informe Dicotómico — WRF-Chem / SINAICA")
    print("=" * 62)
    print(f"  Entrada          : {args.entrada}")
    print(f"  Salida           : {salida}")
    print(f"  Mes              : {args.mes or 'todos'}")
    print(f"  Ciudades         : {ciudades_filtro or 'todas'}")
    print(f"  Categoría NOM-172: {cat_key.upper().replace('_',' ')} "
          f"(riesgo {'Alto/naranja' if cat_key == 'mala' else 'Muy Alto/rojo'})")
    print(f"  Umbrales activos :")
    for cont, val in UMBRALES.items():
        unidad = META_CONT[cont]["unidad"]
        print(f"    {cont:<6} {val:>6.1f} {unidad}")
    print(f"  QC — techo PM2.5 : {LIMITES_VALIDOS['PM25'][1]} µg/m³")
    print(f"  QC — factor IQR  : {IQR_FACTOR}  (k·IQR sobre obs PM2.5)")
    print("=" * 62)

    datos      = leer_csvs(args.entrada, ciudades_filtro, args.mes)
    resultados = calcular_resultados(datos)

    if args.csv_auditoria:
        exportar_csv(resultados, args.csv_auditoria)

    generar_docx(resultados, salida, cat_key)
    print(f"[DONE] {salida}")


if __name__ == "__main__":
    main()
