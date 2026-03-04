#!/usr/bin/env python
# coding: utf-8

# # Proceso de extraccion de datos por ciudad

# In[6]:


#  Para Ozono
import xarray as xr
import pandas as pd
import numpy as np
import os
import calendar
from datetime import datetime

# Configuración
DIR_BASE = "/LUSTRE/OPERATIVO/EXTERNO-salidas/WRF-CHEM/2026/"
CIUDADES = [
    {"nombre": "CDMX", "lat1": 19.20, "lon1": -99.30, "lat2": 19.70, "lon2": -98.85},
    {"nombre": "Toluca", "lat1": 19.23, "lon1": -99.72, "lat2": 19.39, "lon2": -99.50},
    {"nombre": "Puebla", "lat1": 18.95, "lon1": -98.32, "lat2": 19.12, "lon2": -98.10},
    {"nombre": "Tlaxcala", "lat1": 19.29, "lon1": -98.26, "lat2": 19.36, "lon2": -98.15},
    {"nombre": "Pachuca", "lat1": 20.03, "lon1": -98.80, "lat2": 20.13, "lon2": -98.67},
    {"nombre": "Cuernavaca", "lat1": 18.89, "lon1": -99.26, "lat2": 18.98, "lon2": -99.14},
    {"nombre": "SJdelRio", "lat1": 20.36, "lon1": -100.01, "lat2": 20.41, "lon2": -99.93}
]
MES = "02"
ANIO = "2026"

# Determinar el último día del mes
_, ultimo_dia = calendar.monthrange(int(ANIO), int(MES))

# Generar lista de archivos para el mes completo
fechas = pd.date_range(start=f"{ANIO}-{MES}-01", end=f"{ANIO}-{MES}-{ultimo_dia}")
archivos = [os.path.join(DIR_BASE, f"wrfout_d01_{fecha.strftime('%Y-%m-%d')}_00:00:00") 
            for fecha in fechas]

# Diccionario para almacenar resultados por ciudad
resultados_por_ciudad = {ciudad["nombre"]: [] for ciudad in CIUDADES}

def procesar_archivo(archivo, ciudad):
    """Procesa un archivo NetCDF y extrae máximos diarios de O3 para una ciudad"""
    try:
        ds = xr.open_dataset(archivo, decode_times=False)
    except FileNotFoundError:
        print(f"Archivo no encontrado: {archivo}")
        return None
    except Exception as e:
        print(f"Error procesando {archivo}: {e}")
        return None
    
    # Extraer coordenadas
    lat = ds["XLAT"][0, :, :]
    lon = ds["XLONG"][0, :, :]
    
    # Crear máscara para la región
    mask = (lat >= ciudad["lat1"]) & (lat <= ciudad["lat2"]) & \
           (lon >= ciudad["lon1"]) & (lon <= ciudad["lon2"])
    
    # Extraer O3 superficial (primer nivel) en ppbv
    o3_sfc = ds["o3"].isel(bottom_top=0).where(mask) * 1000  # ppmv -> ppbv
    
    # Definir rangos horarios para cada día local (UTC-6)
    rangos = {
        "dia1": slice(6, 30),  # Horas 6 a 29 (24 horas)
        "dia2": slice(30, 54), # Horas 30 a 53 (24 horas)
        "dia3": slice(54, 72)  # Horas 54 a 71 (18 horas)
    }
    
    # Calcular máximos para cada día
    maximos = {}
    for dia, rango in rangos.items():
        o3_dia = o3_sfc.isel(Time=rango)
        if o3_dia.size > 0:
            max_val = o3_dia.max().values.item()
            maximos[dia] = max_val
        else:
            maximos[dia] = np.nan
    
    # Obtener fecha de inicio del pronóstico
    fecha_inicio = archivo.split("_")[-2]
    
    # Cerrar dataset
    ds.close()
    
    return {
        "Fecha": fecha_inicio,
        "Ciudad": ciudad["nombre"],
        **maximos
    }

# Procesar todos los archivos
for archivo in archivos:
    for ciudad in CIUDADES:
        if os.path.exists(archivo):
            resultado = procesar_archivo(archivo, ciudad)
            if resultado:
                resultados_por_ciudad[ciudad["nombre"]].append(resultado)
                print(f"Procesado: {archivo} | Ciudad: {ciudad['nombre']} | "
                      f"Día1: {resultado['dia1']:.2f} ppbv | "
                      f"Día2: {resultado['dia2']:.2f} ppbv | "
                      f"Día3: {resultado['dia3']:.2f} ppbv")
        else:
            print(f"Archivo no encontrado: {archivo}")

# Guardar un archivo CSV por cada ciudad
for nombre_ciudad, resultados in resultados_por_ciudad.items():
    if resultados:
        df = pd.DataFrame(resultados)
        # Reordenar columnas
        df = df[["Fecha", "Ciudad", "dia1", "dia2", "dia3"]]
        df.columns = ["Fecha", "Ciudad", "mod_dia1", "mod_dia2", "mod_dia3"]
        
        # Nombre del archivo de salida
        output_csv = f"maximos_diarios_o3_{nombre_ciudad}_{MES}_{ANIO}.csv"
        df.to_csv(output_csv, index=False)
        print(f"\nResultados para {nombre_ciudad} guardados en: {output_csv}")
        print(f"Total de registros para {nombre_ciudad}: {len(resultados)}")
    else:
        print(f"\nNo se procesaron archivos para {nombre_ciudad}. Verifica la disponibilidad de datos.")


# ## Extrae PM10 y PM2.5

# In[ ]:


import xarray as xr
import pandas as pd
import numpy as np
import os
import calendar
from datetime import datetime

# ==========================================
# CONFIGURACIÓN
# ==========================================
DIR_BASE = "/LUSTRE/OPERATIVO/EXTERNO-salidas/WRF-CHEM/2026/"
CIUDADES = [
    {"nombre": "CDMX", "lat1": 19.20, "lon1": -99.30, "lat2": 19.70, "lon2": -98.85},
    {"nombre": "Toluca", "lat1": 19.23, "lon1": -99.72, "lat2": 19.39, "lon2": -99.50},
    {"nombre": "Puebla", "lat1": 18.95, "lon1": -98.32, "lat2": 19.12, "lon2": -98.10},
    {"nombre": "Tlaxcala", "lat1": 19.29, "lon1": -98.26, "lat2": 19.36, "lon2": -98.15},
    {"nombre": "Pachuca", "lat1": 20.03, "lon1": -98.80, "lat2": 20.13, "lon2": -98.67},
    {"nombre": "Cuernavaca", "lat1": 18.89, "lon1": -99.26, "lat2": 18.98, "lon2": -99.14},
    {"nombre": "SJdelRio", "lat1": 20.36, "lon1": -100.01, "lat2": 20.41, "lon2": -99.93}
]

# Contaminantes a procesar y nombres de variables comunes en WRF-Chem
# Se intentará buscar el nombre exacto, si no, se usan alternativos
POLLUTANTS = {
    "PM10": ["PM10", "pm10"],
    "PM25": ["PM2_5_DRY", "PM2_5", "pm2_5_dry", "pm2_5"] 
}

MES = "02"
ANIO = "2026"

# ==========================================
# FUNCIONES AUXILIARES
# ==========================================

def encontrar_variable(ds, lista_nombres):
    """Busca la primera variable disponible en el dataset desde una lista de posibles nombres."""
    for nombre in lista_nombres:
        if nombre in ds.variables:
            return nombre
    return None

def procesar_archivo_pm(archivo, ciudad, tipo_pm):
    """
    Procesa un archivo NetCDF y extrae el promedio de los máximos espaciales 
    para las ventanas de tiempo definidas.
    tipo_pm: 'PM10' o 'PM25' (para identificar el grupo de variables)
    """
    try:
        # decode_times=False se mantiene para usar índices de hora directos como en el original
        ds = xr.open_dataset(archivo, decode_times=False)
    except FileNotFoundError:
        return None
    except Exception as e:
        print(f"Error abriendo {archivo}: {e}")
        return None
    
    # 1. Identificar la variable correcta en el archivo
    var_nombre = encontrar_variable(ds, POLLUTANTS[tipo_pm])
    
    if var_nombre is None:
        print(f"Advertencia: No se encontró variable PM en {archivo} para {tipo_pm}. Opciones buscadas: {POLLUTANTS[tipo_pm]}")
        ds.close()
        return None
    
    try:
        # 2. Extraer coordenadas
        lat = ds["XLAT"][0, :, :]
        lon = ds["XLONG"][0, :, :]
        
        # 3. Crear máscara para la región de la ciudad
        mask = (lat >= ciudad["lat1"]) & (lat <= ciudad["lat2"]) & \
               (lon >= ciudad["lon1"]) & (lon <= ciudad["lon2"])
        
        # 4. Extraer variable superficial (primer nivel vertical si existe, o promediar verticalmente)
        # Generalmente PM en wrfout es 3D (Time, bottom_top, lat, lon) o 2D.
        var_data = ds[var_nombre]
        
        if "bottom_top" in var_data.dims:
            # Tomar el nivel superficial (índice 0)
            var_sfc = var_data.isel(bottom_top=0)
        else:
            var_sfc = var_data
            
        # Aplicar máscara espacial
        # Donde la máscara es falsa, se pone NaN para que no afecte el máximo
        var_region = var_sfc.where(mask)
        
        # 5. Calcular el MÁXIMO ESPACIAL por cada paso de tiempo
        # Esto reduce las dimensiones lat/lon, dejando solo Time
        max_espacial_por_hora = var_region.max(dim=['south_north', 'west_east'], skipna=True)
        
        # 6. Definir rangos horarios (Índices de tiempo)
        # Día 1: 6 a 30 (24 horas: índices 6, 7, ..., 29)
        # Día 2: 30 a 54 (24 horas: índices 30, 31, ..., 53)
        # Día 3: 54 a 72 (18 horas: índices 54, 55, ..., 71) -> Nota: Archivo de 72h no llega a 24h aquí
        rangos = {
            "dia1": slice(6, 29),  
            "dia2": slice(30, 53), 
            "dia3": slice(54, 72)  
        }
        
        promedios = {}
        for dia, rango in rangos.items():
            datos_dia = max_espacial_por_hora.isel(Time=rango)
            
            if datos_dia.size > 0:
                # Calculamos el promedio de los máximos espaciales en esa ventana
                # Esto equivale al "promedio móvil de 24h" cuando la ventana es exactamente 24h
                val_prom = datos_dia.mean(skipna=True).values.item()
                promedios[dia] = val_prom
            else:
                promedios[dia] = np.nan
        
        # Obtener fecha de inicio del pronóstico del nombre del archivo
        # Formato esperado: wrfout_d01_YYYY-MM-DD_HH:MM:SS
        try:
            fecha_inicio = archivo.split("_")[-2]
        except:
            fecha_inicio = "Desconocida"
            
        ds.close()
        
        return {
            "Fecha": fecha_inicio,
            **promedios
        }
        
    except Exception as e:
        print(f"Error procesando datos PM en {archivo}: {e}")
        ds.close()
        return None

# ==========================================
# PROCESAMIENTO PRINCIPAL
# ==========================================

# Determinar el último día del mes
_, ultimo_dia = calendar.monthrange(int(ANIO), int(MES))

# Generar lista de archivos para el mes completo
fechas = pd.date_range(start=f"{ANIO}-{MES}-01", end=f"{ANIO}-{MES}-{ultimo_dia}")
archivos = [os.path.join(DIR_BASE, f"wrfout_d01_{fecha.strftime('%Y-%m-%d')}_00:00:00") 
            for fecha in fechas]

# Diccionario anidado: Ciudad -> Contaminante -> Lista de resultados
# Estructura: resultados_por_ciudad["CDMX"]["PM10"] = [ {...}, {...} ]
resultados_por_ciudad = {
    ciudad["nombre"]: {tipo: [] for tipo in POLLUTANTS.keys()} 
    for ciudad in CIUDADES
}

print("Iniciando procesamiento de PM10 y PM2.5...")

for archivo in archivos:
    if not os.path.exists(archivo):
        # print(f"Archivo no encontrado: {archivo}") # Silencioso para no saturar consola
        continue
        
    for ciudad in CIUDADES:
        for tipo_pm in POLLUTANTS.keys():
            resultado = procesar_archivo_pm(archivo, ciudad, tipo_pm)
            
            if resultado:
                # Añadir nombre de ciudad al resultado para el DataFrame
                resultado_final = {"Ciudad": ciudad["nombre"], **resultado}
                resultados_por_ciudad[ciudad["nombre"]][tipo_pm].append(resultado_final)
                
                # Feedback limitado para no saturar (solo cada 10 archivos o error)
                # print(f"OK: {ciudad['nombre']} - {tipo_pm} - {resultado['Fecha']}")

# ==========================================
# GUARDAR RESULTADOS
# ==========================================

print("\nGuardando archivos CSV...")

for nombre_ciudad, datos_ciudad in resultados_por_ciudad.items():
    for tipo_pm, lista_resultados in datos_ciudad.items():
        
        if lista_resultados:
            df = pd.DataFrame(lista_resultados)
            
            # Asegurar orden de columnas
            cols_order = ["Fecha", "Ciudad", "dia1", "dia2", "dia3"]
            # Verificar que todas las columnas existan
            cols_existentes = [c for c in cols_order if c in df.columns]
            df = df[cols_existentes]
            
            # Renombrar columnas para claridad
            mapping = {
                "dia1": "prom_24h_dia1",
                "dia2": "prom_24h_dia2",
                "dia3": "prom_24h_dia3"
            }
            df.rename(columns=mapping, inplace=True)
            
            # Definir nombre del contaminante para el archivo
            nombre_archivo_pm = "PM10" if tipo_pm == "PM10" else "PM25"
            
            # Nombre del archivo de salida
            output_csv = f"maximos_diarios_{nombre_archivo_pm}_{nombre_ciudad}_{MES}_{ANIO}.csv"
            
            df.to_csv(output_csv, index=False)
            print(f"Guardado: {output_csv} ({len(df)} registros)")
        else:
            print(f"Sin datos para guardar: {nombre_ciudad} - {tipo_pm}")

print("\nProceso finalizado.")


# ### Concatena las salidas

# In[ ]:


get_ipython().run_cell_magic('bash', '-l', '#!/bin/bash\n# Script: concatenar_por_ciudad.sh (versión multi-contaminante)\n# Descripción: Concatena archivos de máximos diarios para distintos contaminantes\n#              (O3, PM10, PM25) agrupando por ciudad y contaminante.\n#              Los archivos deben tener el formato:\n#              maximos_diarios_<contaminante>_<ciudad>_<mes>_<año>.csv\n# Uso: ./concatenar_por_ciudad.sh [directorio]\n#      Si no se especifica directorio, se usa el directorio actual.\n\n# Directorio de trabajo\nDIR="${1:-.}"\ncd "$DIR" || { echo "Error: No se puede acceder al directorio $DIR"; exit 1; }\nmkdir -p modelo\n# Patrón de búsqueda  \n#PATRON="maximos_diarios_[A-Z]*.csv" #<--- descomentar para PM10 y PM2.5\nPATRON="maximos_diarios_o3_[A-Z]*.csv"\n# Verificar que existan archivos\nif ! ls $PATRON 2>/dev/null; then\n    echo "No se encontraron archivos con el patrón $PATRON en $DIR"\n    exit 1\nfi\n\n# Extraer combinaciones únicas (contaminante + ciudad) a partir de los nombres\n# Formato: maximos_diarios_<cont>_<ciudad>_MM_AAAA.csv\n# Usamos awk para extraer el tercer y cuarto campo (contaminante y ciudad)\ncombinaciones=$(ls $PATRON 2>/dev/null | awk -F\'_\' \'{print $3"_"$4}\' | sort -u)\n\nif [ -z "$combinaciones" ]; then\n    echo "No se pudieron identificar combinaciones de contaminante y ciudad."\n    exit 1\nfi\n\necho "Combinaciones encontradas:"\necho "$combinaciones"\n\n# Procesar cada combinación\nfor combo in $combinaciones; do\n    contaminante=$(echo "$combo" | cut -d\'_\' -f1)\n    ciudad=$(echo "$combo" | cut -d\'_\' -f2)\n    \n    echo "--------------------------------------------------"\n    echo "Procesando: contaminante=$contaminante, ciudad=$ciudad"\n    \n    # Listar archivos de esta combinación y ordenarlos por año y mes\n    # Los campos son: ..._cont_ciudad_mes_año.csv\n    # Usamos sort numérico: año (campo 6) y mes (campo 5)\n    archivos=$(ls maximos_diarios_${contaminante}_${ciudad}_*.csv 2>/dev/null | sort -t \'_\' -k6,6n -k5,5n)\n    \n    if [ -z "$archivos" ]; then\n        echo "  No hay archivos para $contaminante - $ciudad (esto no debería ocurrir)."\n        continue\n    fi\n    \n    # Archivo de salida (sin mes/año)\n    output="modelo/maximos_diarios_${contaminante}_${ciudad}.csv"\n    \n    # Concatenar: el primer archivo con cabecera, los siguientes sin la primera línea\n    primer_archivo=1\n    for archivo in $archivos; do\n        if [ $primer_archivo -eq 1 ]; then\n            cat "$archivo" > "$output"\n            primer_archivo=0\n        else\n            tail -n +2 "$archivo" >> "$output"\n        fi\n        echo "  Añadido: $archivo"\n    done\n    \n    echo "Archivo generado: $output"\ndone\n\necho "--------------------------------------------------"\necho "Proceso completado."\n')


# ## Combina observaciones con modelo

# In[10]:


get_ipython().run_cell_magic('bash', '-l', '#!/bin/bash\n\n# Script: combinar_obs_modelo.sh (versión multi-contaminante)\n# Descripción: Para cada combinación de ciudad y contaminante presente en los archivos de observación,\n#              calcula el máximo diario observado y lo combina con los pronósticos del modelo.\n#              Genera un archivo CSV por ciudad y contaminante en el directorio \'combinado/\'.\n#              Las observaciones están en \'observado/\' con formato: Ciudad_Contaminante_consolidado.csv\n#              Los modelos están en \'modelo/\' con formato: maximos_diarios_Contaminante_Ciudad.csv\n#              Se aplica conversión de ppmv a ppbv solo para O3.\n\n# Directorios\nOBS_DIR="observado"\nMODELO_DIR="modelo"\nOUTPUT_DIR="combinado"\n\nmkdir -p "$OUTPUT_DIR"\n\n# Mapeo de nombres de ciudad: observación -> modelo\ndeclare -A CIUDAD_MAP=(\n    ["San"]="SJdelRio"\n    # Si alguna otra ciudad tiene nombre diferente, agregar aquí\n    # Ejemplo: ["Mexico"]="CDMX"\n)\n\n# Mapeo de nombres de contaminante: observación -> modelo (sin punto para PM2.5)\ndeclare -A CONT_MAP=(\n    ["O3"]="o3"\n    ["PM10"]="PM10"\n    ["PM2.5"]="PM25"\n)\n\n# Factor de conversión para cada contaminante (multiplicar observaciones)\n# O3: ppmv -> ppbv (x1000). PM10 y PM2.5 se asumen en mismas unidades (µg/m³) sin conversión.\ndeclare -A FACTOR=(\n    ["O3"]=1000\n    ["PM10"]=1\n    ["PM2.5"]=1\n)\n\n# Procesar cada archivo de observación\nfor obs_file in "$OBS_DIR"/*_consolidado.csv; do\n    [ -f "$obs_file" ] || continue  # si no hay archivos, sale\n\n    # Extraer ciudad y contaminante del nombre (ej: Cuernavaca_O3_consolidado.csv)\n    basename=$(basename "$obs_file" _consolidado.csv)   # quita sufijo\n    # basename tiene formato "Ciudad_Contaminante"\n    ciudad_obs=$(echo "$basename" | cut -d\'_\' -f1)\n    contaminante_obs=$(echo "$basename" | cut -d\'_\' -f2)\n\n    # Mapear a nombres de modelo\n    ciudad_modelo="${CIUDAD_MAP[$ciudad_obs]:-$ciudad_obs}"   # si no está en mapa, usar el mismo\n    contaminante_modelo="${CONT_MAP[$contaminante_obs]}"\n    factor="${FACTOR[$contaminante_obs]}"\n\n    if [ -z "$contaminante_modelo" ]; then\n        echo "ERROR: Contaminante \'$contaminante_obs\' no reconocido en $obs_file"\n        continue\n    fi\n\n    echo "=================================================="\n    echo "Procesando: $ciudad_obs ($contaminante_obs) -> modelo: $ciudad_modelo ($contaminante_modelo)"\n\n    # Archivo de modelo esperado\n    modelo_file="$MODELO_DIR/maximos_diarios_${contaminante_modelo}_${ciudad_modelo}.csv"\n\n    # 1. Calcular máximo diario de observaciones (con posible conversión)\n    echo "  Calculando máximo diario de observaciones..."\n    # Usamos awk para leer el archivo de observaciones (formato: date,hour,est1,est2,...)\n    # Se calcula el máximo por fecha, ignorando NA y aplicando factor.\n    awk -F, -v factor="$factor" \'\n    BEGIN { OFS="," }\n    NR>1 {\n        fecha = $1\n        if (!(fecha in max)) max[fecha] = -1\n        for (i=3; i<=NF; i++) {\n            if ($i != "NA" && $i != "") {\n                val = $i * factor\n                if (val > max[fecha]) max[fecha] = val\n            }\n        }\n    }\n    END {\n        for (f in max) {\n            if (max[f] == -1)\n                printf "%s,NA\\n", f\n            else\n                printf "%s,%.3f\\n", f, max[f]\n        }\n    }\' "$obs_file" | sort -t, -k1 > obs_max_$$.tmp\n\n    # 2. Preparar archivo de modelo (si existe)\n    if [ -f "$modelo_file" ]; then\n        echo "  Procesando modelo..."\n        # Extraer Fecha, mod_dia1, dia2, dia3 (columnas 1,3,4,5)\n        # Nota: aunque el contaminante sea PM, las columnas se llaman igual.\n        awk -F, \'NR>1 {print $1","$3","$4","$5}\' "$modelo_file" | sort -t, -k1 > modelo_$$.tmp\n        tiene_modelo=1\n    else\n        echo "  AVISO: No se encuentra $modelo_file. Se usarán NA para pronósticos."\n        tiene_modelo=0\n    fi\n\n    # 3. Combinar observaciones con modelo (left join por fecha)\n    echo "  Combinando datos..."\n    if [ $tiene_modelo -eq 1 ]; then\n        # Usamos awk para hacer left join: cargar modelo en array, luego leer obs\n        awk -F, \'\n        BEGIN { OFS="," }\n        NR==FNR {\n            # modelo: fecha, dia1, dia2, dia3\n            modelo[$1] = $2","$3","$4\n            next\n        }\n        {\n            fecha = $1\n            max_obs = $2\n            if (fecha in modelo)\n                print fecha, max_obs, modelo[fecha]\n            else\n                print fecha, max_obs, "NA,NA,NA"\n        }\' modelo_$$.tmp obs_max_$$.tmp | sort -t, -k1 > combinado_$$.tmp\n    else\n        # Sin modelo: añadir NA en las columnas de pronóstico\n        awk -F, \'{print $1","$2",NA,NA,NA"}\' obs_max_$$.tmp | sort -t, -k1 > combinado_$$.tmp\n    fi\n\n    # 4. Escribir archivo final con encabezado\n    output_file="$OUTPUT_DIR/combinado_${ciudad_obs}_${contaminante_obs}.csv"\n    echo "Fecha,max_obs,mod_dia1,mod_dia2,mod_dia3" > "$output_file"\n    cat combinado_$$.tmp >> "$output_file"\n\n    # Limpiar temporales\n    rm -f obs_max_$$.tmp modelo_$$.tmp combinado_$$.tmp\n\n    echo "  Archivo generado: $output_file"\ndone\n\necho "=================================================="\necho "Proceso completado. Los archivos están en $OUTPUT_DIR"\n')


# In[ ]:


get_ipython().run_cell_magic('bash', '-l', '# Script: ls o.sh\n# Descripción: Para cada ciudad, calcula el máximo diario de ozono observado (ppbv)\n#              y lo combina con los pronósticos del modelo (día1, día2, día3).\n#              Genera un archivo CSV por ciudad con columnas:\n#              Fecha, max_obs, mod_dia1, mod_dia2, mod_dia3.\n\n# Directorios de entrada y salida\nOBS_DIR="observado"\nMODELO_DIR="modelo"\nOUTPUT_DIR="combinado"\n\nmkdir -p "$OUTPUT_DIR"\n\n# Mapeo: nombre en modelo -> nombre base en archivo de observaciones\n# (ej: SJdelRio se llama "San" en observado)\ndeclare -A CIUDAD_MAP=(\n    ["Pachuca"]="Pachuca"\n    ["Puebla"]="Puebla"\n    ["Tlaxcala"]="Tlaxcala"\n    ["Toluca"]="Toluca"\n    ["Cuernavaca"]="Cuernavaca"\n    ["SJdelRio"]="San"\n)\n\n# Procesar cada ciudad\nfor ciudad in "${!CIUDAD_MAP[@]}"; do\n    obs_name="${CIUDAD_MAP[$ciudad]}"\n    obs_file="$OBS_DIR/${obs_name}_O3_consolidado.csv"\n    modelo_file="$MODELO_DIR/maximos_diarios_o3_${ciudad}.csv"\n    output_file="$OUTPUT_DIR/combinado_${ciudad}.csv"\n\n    echo "=================================================="\n    echo "Procesando: $ciudad"\n\n    # Verificar existencia de archivo de observaciones\n    if [ ! -f "$obs_file" ]; then\n        echo "  ERROR: No se encuentra $obs_file"\n        continue\n    fi\n\n    # 1. Calcular máximo diario de observaciones (ppbv)\n    echo "  Calculando máximo diario de observaciones..."\n    awk -F, \'\n    BEGIN { OFS="," }\n    NR>1 {\n        fecha = $1\n        if (!(fecha in max)) max[fecha] = -1\n        for (i=3; i<=NF; i++) {\n            if ($i != "NA" && $i != "") {\n                val = $i * 1000          # convertir ppmv a ppbv\n                if (val > max[fecha]) max[fecha] = val\n            }\n        }\n    }\n    END {\n        for (f in max) {\n            if (max[f] == -1)\n                printf "%s,NA\\n", f\n            else\n                printf "%s,%.3f\\n", f, max[f]\n        }\n    }\' "$obs_file" | sort -t, -k1 > obs_max_$$.tmp\n\n    # 2. Preparar archivo de modelo (si existe)\n    if [ -f "$modelo_file" ]; then\n        echo "  Procesando modelo..."\n        # Extraer Fecha, mod_dia1, dia2, dia3 (columnas 1,3,4,5)\n        awk -F, \'NR>1 {print $1","$3","$4","$5}\' "$modelo_file" | sort -t, -k1 > modelo_$$.tmp\n        tiene_modelo=1\n    else\n        echo "  AVISO: No se encuentra $modelo_file. Se usarán NA para pronósticos."\n        tiene_modelo=0\n    fi\n\n    # 3. Combinar observaciones con modelo (left join por fecha)\n    echo "  Combinando datos..."\n    if [ $tiene_modelo -eq 1 ]; then\n        # Cargar modelo en memoria y hacer join con observaciones\n        awk -F, \'\n        BEGIN { OFS="," }\n        NR==FNR {\n            # modelo: fecha, dia1, dia2, dia3\n            modelo[$1] = $2","$3","$4\n            next\n        }\n        {\n            fecha = $1\n            max_obs = $2\n            if (fecha in modelo)\n                print fecha, max_obs, modelo[fecha]\n            else\n                print fecha, max_obs, "NA,NA,NA"\n        }\' modelo_$$.tmp obs_max_$$.tmp | sort -t, -k1 > combinado_$$.tmp\n    else\n        # Sin modelo: añadir NA en las columnas de pronóstico\n        awk -F, \'{print $1","$2",NA,NA,NA"}\' obs_max_$$.tmp | sort -t, -k1 > combinado_$$.tmp\n    fi\n\n    # 4. Escribir archivo final con encabezado\n#    echo "Fecha,max_obs,mod_dia1,mod_dia2,mod_dia3" > "$output_file"\n    echo "Fecha,obs,mod_dia0,mod_dia1,mod_dia2" > "$output_file"\n    cat combinado_$$.tmp >> "$output_file"\n\n    # Limpiar archivos temporales\n    rm -f obs_max_$$.tmp modelo_$$.tmp combinado_$$.tmp\n\n    echo "  Archivo generado: $output_file"\ndone\n\necho "=================================================="\necho "Proceso completado. Los archivos están en $OUTPUT_DIR"\n')


# ### Elimina filas con NA,NA,NA

# In[11]:


get_ipython().run_cell_magic('bash', '-l', '#!/bin/bash\n\n# Script: limpiar_combinados.sh\n# Descripción: Elimina filas donde los tres pronósticos del modelo son NA.\n# Uso: ./limpiar_combinados.sh [directorio]\n# Si no se especifica directorio, se usa el directorio actual.\n\nDIR="${1:-.}"\ncd combinado\n# Buscar todos los archivos CSV combinados (ajusta el patrón si es necesario)\nfor archivo in combinado_*.csv; do\n    [ -e "$archivo" ] || continue  # si no hay archivos, sale\n    \n    echo "Procesando: $archivo"\n    \n    # Archivo temporal\n    temp_file="${archivo}.tmp"\n    \n    # Conservar encabezado y filtrar filas que NO terminen en ",NA,NA,NA"\n    head -1 "$archivo" > "$temp_file"\n    tail -n +2 "$archivo" | grep -v \',NA,NA,NA\' >> "$temp_file"\n    \n    # Reemplazar original por el limpio (opcional: puedes renombrar o mover)\n    mv "$temp_file" "$archivo"\n    \n    echo "  Limpiado: $(basename "$archivo")"\ndone\ncd -\necho "Proceso completado."\n')


# ##  En combinados desplaza los registros

# In[14]:


get_ipython().run_cell_magic('bash', '-l', '#!/bin/bash\n\n# ==========================================================\n# Ajusta horizontes de pronóstico para múltiples archivos\n# Patron: combinado_*.csv\n# Desplaza:\n#   mod_dia2 → +1 día\n#   mod_dia3 → +2 días\n# Compatible macOS y GNU/Linux\n# ==========================================================\n\ncd /home/agustin/FQA/agustin/DOMAINS/evaluacion/combinado\nmkdir -p ajustados\n# Detectar tipo de date\nif date --version >/dev/null 2>&1; then\n    DATE_CMD="gnu"\nelse\n    DATE_CMD="bsd"\nfi\n\nsumar_dias() {\n    fecha=$1\n    dias=$2\n\n    if [ "$DATE_CMD" = "gnu" ]; then\n        date -d "$fecha +$dias day" +%Y-%m-%d\n    else\n        date -j -f %Y-%m-%d "$fecha" -v+${dias}d +%Y-%m-%d\n    fi\n}\n\nfor input in combinado_*.csv; do\n\n    echo "Procesando: $input"\n\n    tmpfile=$(mktemp)\n\n{\n    awk -F\',\' \'\n    BEGIN{ OFS="," }\n\n    NR==1{ next }\n\n    {\n        fecha=$1\n        obs[fecha]=$2\n        dia1[fecha]=$3\n\n        cmd="date -d \\"" fecha " +1 day\\" +%Y-%m-%d 2>/dev/null"\n        cmd | getline f2\n        close(cmd)\n\n        if (f2=="") {\n            cmd="date -j -f %Y-%m-%d " fecha " -v+1d +%Y-%m-%d"\n            cmd | getline f2\n            close(cmd)\n        }\n\n        cmd="date -d \\"" fecha " +2 day\\" +%Y-%m-%d 2>/dev/null"\n        cmd | getline f3\n        close(cmd)\n\n        if (f3=="") {\n            cmd="date -j -f %Y-%m-%d " fecha " -v+2d +%Y-%m-%d"\n            cmd | getline f3\n            close(cmd)\n        }\n\n        dia2[f2]=$4\n        dia3[f3]=$5\n        fechas[fecha]=1\n        fechas[f2]=1\n        fechas[f3]=1\n    }\n\n    END{\n        for (f in fechas)\n            printf "%s,%s,%s,%s,%s\\n",\n            f,\n            (f in obs ? obs[f] : "NA"),\n            (f in dia1 ? dia1[f] : "NA"),\n            (f in dia2 ? dia2[f] : "NA"),\n            (f in dia3 ? dia3[f] : "NA")\n    }\n    \' "$input" | sort -t\',\' -k1,1\n} > tmp_sorted.csv\n\n# Insertar encabezado correctamente\n{\n    echo "Fecha,max_obs,mod_dia1,mod_dia2,mod_dia3"\n    cat tmp_sorted.csv\n} > "ajustados/${input}"\n\nrm tmp_sorted.csv\n    rm "$tmpfile"\n\n    echo "  -> Generado: ajustados/${input}"\n\ndone\ncd -\necho "Proceso terminado."\n')


# #  Calculo de estadisticos
# 

# In[20]:


import os
import glob
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches

# ============================
# Funciones continuas (bootstrap)
# ============================

def bootstrap_continuous(obs, mod, B=10000):
    """
    Bootstrap de MAE, MSE, RMSE, BIAS y R^2
    """
    n = len(obs)
    
    mae_b, mse_b, rmse_b, bias_b, r2_b = [], [], [], [], []
    
    obs_mean = np.mean(obs)
    sst = np.sum((obs - obs_mean) ** 2)
    
    for _ in range(B):
        idx = np.random.choice(n, n, replace=True)
        o, m = obs[idx], mod[idx]
        
        mae = np.mean(np.abs(m - o))
        mse = np.mean((m - o) ** 2)
        rmse = np.sqrt(mse)
        bias = np.mean(m - o)
        
        ssr = np.sum((o - m) ** 2)
        r2 = 1.0 - ssr / sst if sst > 0 else np.nan
        
        mae_b.append(mae)
        mse_b.append(mse)
        rmse_b.append(rmse)
        bias_b.append(bias)
        r2_b.append(r2)
    
    def resumen(arr):
        return {
            "mean": np.nanmean(arr),
            "std": np.nanstd(arr),
            "ci95": (
                np.nanpercentile(arr, 2.5),
                np.nanpercentile(arr, 97.5)
            )
        }
    
    return {
        "MAE": resumen(mae_b),
        "MSE": resumen(mse_b),
        "RMSE": resumen(rmse_b),
        "BIAS": resumen(bias_b),
        "R2": resumen(r2_b)
    }

# ============================
# Métricas dicotómicas
# ============================

def categorical_once(obs, mod, threshold=135):
    obs_bin = (obs > threshold).astype(int)
    mod_bin = (mod > threshold).astype(int)
    
    H = np.sum((obs_bin == 1) & (mod_bin == 1))  # hits
    M = np.sum((obs_bin == 1) & (mod_bin == 0))  # misses
    F = np.sum((obs_bin == 0) & (mod_bin == 1))  # false alarms
    C = np.sum((obs_bin == 0) & (mod_bin == 0))  # correct negatives
    
    N = H + M + F + C
    
    def sd(a, b): return a / b if b > 0 else np.nan
    
    PC   = sd(H + C, N)
    POD  = sd(H, H + M)
    FAR  = sd(F, H + F)
    SR   = sd(H, H + F)
    POFD = sd(F, F + C)
    CSI  = sd(H, H + M + F)
    TSS  = POD - POFD
    
    HSS = sd(
        2 * (H*C - F*M),
        (H+M)*(M+C) + (H+F)*(F+C)
    )
    
    GS = sd(
        H - (H+M)*(H+F)/N,
        H + M + F - (H+M)*(H+F)/N
    )
    
    OR   = sd(H*C, F*M)
    ORSS = sd(OR - 1, OR + 1)
    
    CRF = sd((H+M)*(H+F) + (F+C)*(M+C), N)
    CH  = sd((H+M)*(H+F), N)
    
    return {
        "PC": PC, "POD": POD, "FAR": FAR, "SR": SR,
        "TSS": TSS, "CSI": CSI, "HSS": HSS, "GS": GS,
        "POFD": POFD, "OR": OR, "ORSS": ORSS,
        "CRF": CRF, "CH": CH
    }

def bootstrap_categorical(obs, mod, B=10000, threshold=135):
    n = len(obs)
    keys = categorical_once(obs, mod, threshold).keys()
    store = {k: [] for k in keys}
    
    for _ in range(B):
        idx = np.random.choice(n, n, replace=True)
        res = categorical_once(obs[idx], mod[idx], threshold)
        for k in keys:
            store[k].append(res[k])
    
    def resumen(arr):
        return {
            "mean": np.nanmean(arr),
            "std": np.nanstd(arr),
            "ci95": (
                np.nanpercentile(arr, 2.5),
                np.nanpercentile(arr, 97.5)
            )
        }
    
    return {k: resumen(v) for k, v in store.items()}

# ============================
# Clase para diagrama de Taylor
# ============================

class TaylorDiagram:
    def __init__(self, ref_std, fig=None, label='Observaciones'):
        self.ref_std = ref_std

        if fig is None:
            fig = plt.figure(figsize=(8, 7))

        self.ax = fig.add_subplot(111, polar=True)

        # Configuración estándar
        self.ax.set_theta_zero_location("E")   # correlación=1 en eje X
        self.ax.set_theta_direction(1)
        self.ax.set_thetamin(0)
        self.ax.set_thetamax(90)

        # Escala radial (Desviación estándar)
        self.max_std = ref_std * 1.6
        self.ax.set_rlim(0, self.max_std)

        # --- Etiquetas de correlación ---
        corr_ticks = np.array([0, 0.2, 0.4, 0.6, 0.8, 0.9, 0.95, 1])
        angles = np.degrees(np.arccos(corr_ticks))
        self.ax.set_thetagrids(angles,
                               labels=[f"{c:.2f}" for c in corr_ticks])

        # Etiqueta radial
        self.ax.set_ylabel("Desviación estándar (ppb)", labelpad=25)

        # Etiqueta angular
        self.ax.text(np.pi/4, self.max_std * 1.15,
                     "Correlación (r)",
                     horizontalalignment='center')

        # Punto de referencia
        self.ax.plot(0, ref_std, 'ko', markersize=9, label=label)

    def add_sample(self, std, corrcoef, label, marker='o'):
        corrcoef = np.clip(corrcoef, -1, 1)
        theta = np.arccos(corrcoef)

        self.ax.plot(theta, std,
                     marker=marker,
                     linestyle='None',
                     markersize=8,
                     label=label)

    def add_rmse_contours(self, levels=5):
        rs, ts = np.meshgrid(
            np.linspace(0, self.max_std, 200),
            np.linspace(0, np.pi/2, 200)
        )

        rmse = np.sqrt(
            self.ref_std**2 +
            rs**2 -
            2 * self.ref_std * rs * np.cos(ts)
        )

        contours = self.ax.contour(ts, rs, rmse,
                                   levels=levels,
                                   colors='0.6')

        self.ax.clabel(contours, inline=True, fontsize=8)

        # Etiqueta RMSE
        self.ax.text(np.pi/3, self.max_std * 0.6,
                     "RMSE (ppb)",
                     rotation=35)

        return contours

# ============================
# Funciones de graficado adicionales
# ============================

def crear_scatter_plot(obs, mod, titulo, variable, output_path):
    """
    Crea un scatter plot observado vs pronosticado con línea 1:1.
    Incluye el coeficiente de correlación y RMSE en el gráfico.
    """
    fig, ax = plt.subplots(figsize=(6, 6))
    
    # Calcular métricas
    mask = ~(np.isnan(obs) | np.isnan(mod))
    o = obs[mask]
    m = mod[mask]
    if len(o) == 0:
        plt.close()
        return None
    
    corr = np.corrcoef(o, m)[0, 1]
    rmse = np.sqrt(np.mean((m - o) ** 2))
    bias = np.mean(m - o)
    
    ax.scatter(o, m, alpha=0.6, edgecolors='k', linewidth=0.5)
    
    # Línea 1:1
    min_val = min(o.min(), m.min())
    max_val = max(o.max(), m.max())
    ax.plot([min_val, max_val], [min_val, max_val], 'r--', lw=1, label='1:1')
    
    ax.set_xlabel(f'Observado ({variable})')
    ax.set_ylabel(f'Pronóstico ({variable})')
    ax.set_title(titulo)
    ax.grid(True, linestyle=':', alpha=0.7)
    
    # Añadir texto con métricas
    textstr = f'n = {len(o)}\nR = {corr:.3f}\nRMSE = {rmse:.2f}\nBIAS = {bias:.2f}'
    props = dict(boxstyle='round', facecolor='wheat', alpha=0.5)
    ax.text(0.05, 0.95, textstr, transform=ax.transAxes, fontsize=10,
            verticalalignment='top', bbox=props)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    return output_path

def crear_series_tiempo(df, variable, output_path):
    """
    Crea un gráfico de series de tiempo con observado y los tres pronósticos.
    df debe contener columnas: 'Fecha', 'max_obs', 'mod_dia1', 'mod_dia2', 'mod_dia3'
    (los nombres de los pronósticos se renombrarán en la leyenda).
    """
    fig, ax = plt.subplots(figsize=(12, 5))
    
    # Convertir Fecha a datetime si no lo está
    if not pd.api.types.is_datetime64_any_dtype(df['Fecha']):
        df['Fecha'] = pd.to_datetime(df['Fecha'])
    
    # Graficar
    ax.plot(df['Fecha'], df['max_obs'], 'k-', label='Observado', linewidth=1.5)
    ax.plot(df['Fecha'], df['mod_dia1'], 'r-', label='Día 0', linewidth=1, alpha=0.7)
    ax.plot(df['Fecha'], df['mod_dia2'], 'g-', label='Día 1', linewidth=1, alpha=0.7)
    ax.plot(df['Fecha'], df['mod_dia3'], 'b-', label='Día 2', linewidth=1, alpha=0.7)
    
    ax.set_xlabel('Fecha')
    ax.set_ylabel(f'Concentración ({variable})')
    ax.set_title('Serie temporal: Observado vs Pronósticos')
    ax.legend(loc='best')
    ax.grid(True, linestyle=':', alpha=0.7)
    
    # Formato del eje x para fechas
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    plt.xticks(rotation=45)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    return output_path

# ============================
# Procesamiento de todas las ciudades
# ============================

def detectar_variable(nombre_archivo):
    """
    Extrae la variable (O3, PM10, PM2.5) del nombre del archivo.
    Se asume formato: combinado_CIUDAD_VARIABLE.csv
    Si no se encuentra, retorna 'O3' por defecto.
    """
    base = os.path.basename(nombre_archivo).replace(".csv", "")
    if base.startswith("combinado_"):
        resto = base[10:]  # quitar "combinado_"
    else:
        resto = base
    
    partes = resto.split("_")
    if len(partes) >= 2 and partes[-1] in ["O3", "PM10", "PM2.5"]:
        return partes[-1]
    return "O3"

def procesar_ciudad(archivo):
    """
    Procesa un archivo combinado de una ciudad y genera el documento Word.
    """
    # Extraer nombre de la ciudad y variable
    nombre_base = os.path.basename(archivo).replace(".csv", "")
    # Eliminar prefijo "combinado_" si existe
    if nombre_base.startswith("combinado_"):
        resto = nombre_base[10:]  # quitar "combinado_"
    else:
        resto = nombre_base
    
    # Intentar separar ciudad y variable
    partes = resto.split("_")
    if len(partes) >= 2 and partes[-1] in ["O3", "PM10", "PM2.5"]:
        variable = partes[-1]
        ciudad = "_".join(partes[:-1])
    else:
        variable = "O3"
        ciudad = resto
    
    print(f"\n--- Procesando ciudad: {ciudad}, variable: {variable} ---")

    # Leer datos
    df = pd.read_csv(archivo)
    
    # Verificar que existan las columnas necesarias max_obs,mod_dia1,mod_dia2,mod_dia3
    columnas_requeridas = ["max_obs", "mod_dia1", "mod_dia2", "mod_dia3"]
    for col in columnas_requeridas:
        if col not in df.columns:
            print(f"  Error: Columna '{col}' no encontrada en {archivo}. Se omite.")
            return
    
    # Guardar una copia con fechas para series de tiempo
    df_fechas = df.copy()
    
    # Filtrar filas con NaN en las columnas de interés para bootstrap
    df_clean = df[columnas_requeridas].dropna()
    n_total = len(df)
    n_validos = len(df_clean)
    
    if n_validos == 0:
        print(f"  No hay datos válidos para {ciudad} (todos NA). Se omite.")
        return
    
    obs = df_clean["max_obs"].values
    fecha_min = df["Fecha"].min() if "Fecha" in df.columns else "desconocida"
    fecha_max = df["Fecha"].max() if "Fecha" in df.columns else "desconocida"
    
    # Diccionario para guardar resultados por horizonte
    resultados = {}
    
    # Calcular bootstrap para cada horizonte
    for horizonte in ["mod_dia1", "mod_dia2", "mod_dia3"]:
        mod = df_clean[horizonte].values
        
        # Bootstrap continuo
        stats_cont = bootstrap_continuous(obs, mod, B=10000)
        # Convertir a tuplas (mean, std, lo, hi) para facilitar exportación
        cont_tuplas = {m: (v['mean'], v['std'], v['ci95'][0], v['ci95'][1]) 
                       for m, v in stats_cont.items()}
        
        # Bootstrap categórico
        stats_cat = bootstrap_categorical(obs, mod, B=10000, threshold=135)
        cat_tuplas = {m: (v['mean'], v['std'], v['ci95'][0], v['ci95'][1]) 
                      for m, v in stats_cat.items()}
        
        resultados[horizonte] = {
            "continuos": cont_tuplas,
            "dicotomicos": cat_tuplas
        }
    
    # --- Diagrama de Taylor ---
    std_obs = np.std(obs, ddof=1)
    fig_taylor = plt.figure(figsize=(8, 7))
    taylor = TaylorDiagram(std_obs, fig=fig_taylor, label='Observaciones')
    taylor.add_rmse_contours(levels=5)
    
    colores = {'mod_dia1': 'r', 'mod_dia2': 'g', 'mod_dia3': 'b'}
    marcadores = {'mod_dia1': 'o', 'mod_dia2': 's', 'mod_dia3': '^'}
    nombres_horizontes = {'mod_dia1': 'Día 0', 'mod_dia2': 'Día 1', 'mod_dia3': 'Día 2'}
    
    for horizonte in ["mod_dia1", "mod_dia2", "mod_dia3"]:
        mod = df_clean[horizonte].values
        std_mod = np.std(mod, ddof=1)
        corr = np.corrcoef(obs, mod)[0, 1]
        taylor.add_sample(std_mod, corr, 
                          label=nombres_horizontes[horizonte], 
                          marker=marcadores[horizonte])
    
    plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.0))
    taylor_path = f"/tmp/taylor_{ciudad}_{variable}.png"
    plt.savefig(taylor_path, dpi=150, bbox_inches='tight')
    plt.close(fig_taylor)
    
    # --- Scatter plots para cada horizonte ---
    scatter_paths = {}
    for horizonte in ["mod_dia1", "mod_dia2", "mod_dia3"]:
        mod = df_clean[horizonte].values
        scatter_path = f"/tmp/scatter_{ciudad}_{variable}_{horizonte}.png"
        titulo = f"{ciudad} - {nombres_horizontes[horizonte]}"
        ruta = crear_scatter_plot(obs, mod, titulo, variable, scatter_path)
        if ruta:
            scatter_paths[horizonte] = ruta
    
    # --- Serie de tiempo (usando df_fechas con posibles NA, pero se grafican con huecos) ---
    time_path = f"/tmp/timeseries_{ciudad}_{variable}.png"
    crear_series_tiempo(df_fechas, variable, time_path)
    
    # --- Crear documento Word ---
    doc = Document()
    titulo = f"Reporte de Validación: Modelo WRF-Chem vs RAMA - {ciudad} ({variable})"
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(f"Periodo evaluado: {fecha_min} a {fecha_max}. "
                      f"Días con datos válidos (para bootstrap): {n_validos} de {n_total}. "
                      f"Métricas obtenidas vía Bootstrap (B=10000, IC 95%).")
    
    # Insertar diagrama de Taylor
    doc.add_heading("Diagrama de Taylor", level=2)
    doc.add_picture(taylor_path, width=Inches(6))
    
    # Insertar serie de tiempo
    doc.add_heading("Serie Temporal", level=2)
    doc.add_picture(time_path, width=Inches(6))
    
    # Insertar scatter plots (uno por horizonte)
    doc.add_heading("Gráficos de Dispersión", level=2)
    for horizonte in ["mod_dia1", "mod_dia2", "mod_dia3"]:
        if horizonte in scatter_paths:
            doc.add_heading(nombres_horizontes[horizonte], level=3)
            doc.add_picture(scatter_paths[horizonte], width=Inches(5))
    
    # Tablas por horizonte
    for horizonte_orig, nombres_mostrar in [("mod_dia1", "Día 0"), 
                                             ("mod_dia2", "Día 1"), 
                                             ("mod_dia3", "Día 2")]:
        doc.add_heading(f"Horizonte de Pronóstico: {nombres_mostrar}", level=2)
        bloques = resultados[horizonte_orig]
        
        for tipo, titulo_tabla in [("continuos", "Métricas Continuas"), 
                                    ("dicotomicos", "Métricas Dicotómicas (>135 ppb)")]:
            doc.add_heading(titulo_tabla, level=3)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = \
                "Métrica", "Media", "Desv. Std", "IC 2.5%", "IC 97.5%"
            
            for met, (mean, std, lo, hi) in bloques[tipo].items():
                row = table.add_row().cells
                row[0].text = met
                row[1].text = f"{mean:.3f}"
                row[2].text = f"{std:.3f}"
                row[3].text = f"{lo:.3f}"
                row[4].text = f"{hi:.3f}"
    
    # Guardar documento
    output_docx = f"evaluacion_{variable}_{ciudad}.docx"
    doc.save(output_docx)
    print(f"  Documento guardado: {output_docx}")
    
    # Eliminar imágenes temporales
    os.remove(taylor_path)
    os.remove(time_path)
    for path in scatter_paths.values():
        os.remove(path)


# ============================
# Ejecución principal
# ============================

if __name__ == "__main__":
    # Buscar todos los archivos combinados
    archivos = glob.glob("combinado/ajustados/combinado_*.csv")
    
    if not archivos:
        print("No se encontraron archivos en 'combinado/combinado_*.csv'")
    else:
        for archivo in archivos:
            procesar_ciudad(archivo)
    
    print("\nProcesamiento completado.")


# In[ ]:




